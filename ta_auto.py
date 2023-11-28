
import os

#설치 완료후 환경설정
# python 패키지로 JAVA_HOME 설정하기
os.environ["JAVA_HOME"] = "/opt/conda"

# 필요 패키지 추가
import time
import datetime
import pickle
import docx
from docx import Document
import re
import sys, io
import tempfile
import shutil
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Border, Side, PatternFill, Alignment, Protection, GradientFill, Color
from openpyxl.cell import MergedCell
from openpyxl import load_workbook
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION

# 한글 토큰화 konlpy 모듈
# from konlpy.tag import Okt
from ckonlpy.tag import Twitter
# okt = Okt()
twitter = Twitter() # twitter가 okt보다 성능 높음

# from gensim.models import Word2Vec

# 문서작업 중 필요 모듈
import re
import glob
import warnings
# import gensim

# model checkpoint
import mygsmod as gs
from mygsmod import ModelCheckpoint

import pandas as pd
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import seaborn as sns
import tempfile

import streamlit as st

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics import confusion_matrix

import torch.nn as nn
import torch
import torch.optim as optim
import torchsummary
from torchsummary import summary
from torch.nn import functional as F
from torch.utils.data import Dataset, DataLoader
from torchtext.vocab import build_vocab_from_iterator
from torch.nn.utils.rnn import pad_sequence

# 문서 style copy 함수
def copy_style(src, dest):
    if isinstance(src, Font):
        dest.font = Font(name=src.name, size=src.size, bold=src.bold, italic=src.italic,
                         vertAlign=src.vertAlign, underline=src.underline, strike=src.strike,
                         color=src.color)
    elif isinstance(src, Border):
        dest.border = src
    elif isinstance(src, (PatternFill, GradientFill)):
        dest.fill = src
    elif isinstance(src, Alignment):
        dest.alignment = src
    elif isinstance(src, Protection):
        dest.protection = src
        
# worklist sheet 별로 문서 분리 함수 
def split_excel_sheets_to_files(input_file):
    # Determine file type (Excel or CSV)
    if input_file.name.endswith('.xlsx'):
        wb = openpyxl.load_workbook(input_file)
        file_type = 'xlsx'
    elif input_file.name.endswith('.csv'):
        df = pd.read_csv(input_file)
        wb = openpyxl.Workbook()
        ws = wb.active
        for r, row in enumerate(df.values, 1):
            for c, value in enumerate(row, 1):
                ws.cell(row=r, column=c, value=value)
        file_type = 'csv'
    else:
        st.error("Unsupported file format. Please upload an Excel or CSV file.")
        return
    
    # Process and return files
    output_files = []
    for sheetname in wb.sheetnames:
        # Create a new workbook with only one sheet
        new_wb = openpyxl.Workbook()
        new_ws = new_wb.active
        
        # Set the name of the new sheet to the original sheet's name
        new_ws.title = sheetname

        # Copy data from the old sheet to the new one
        for row in wb[sheetname].iter_rows():
            for cell in row:
                new_cell = new_ws[cell.coordinate]
                new_cell.value = cell.value
                # Copy styles using the helper function
                copy_style(cell.font, new_cell)
                copy_style(cell.border, new_cell)
                copy_style(cell.fill, new_cell)
                new_cell.number_format = cell.number_format
                copy_style(cell.alignment, new_cell)
                copy_style(cell.protection, new_cell)
             
        # Save to a BytesIO object
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            new_wb.save(tmp.name)
            tmp.seek(0)
            with open(tmp.name, 'rb') as file_data:
                output_files.append((f"{sheetname}.xlsx", file_data.read()))  

    return output_files

# worklist 형태 변경 함수
def worklist_type_transform(uploaded_file):
    # Read the file based on its format# Determine file type (Excel or CSV)
    if uploaded_file.type == "text/csv":
        df = pd.read_csv(uploaded_file)
    else:  # Assuming Excel
        df = pd.read_excel(uploaded_file)

    # Perform the required transformations
    df.drop(df.columns[[0, 3, 5, 6]], axis=1, inplace=True)
    df = df.iloc[4:].reset_index(drop=True)
    condition = df[df.columns[1]].apply(lambda x: '장치기술1팀' in str(x) or pd.isna(x))
    df_filtered = df[condition]
    
    # Save DataFrame to a temporary file
    _, temp_file_path = tempfile.mkstemp(suffix='.xlsx')
    with pd.ExcelWriter(temp_file_path, engine='openpyxl') as writer:
        df_filtered.to_excel(writer, index=False)

    return temp_file_path

# 형태 변경된 worklist 전처리
def insert_row_based_on_condition(file_path, new_file_path):
    # Load Excel file
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active

    # Get the initial count of rows
    initial_rows = list(sheet.iter_rows(values_only=True))  # Create a snapshot of initial rows

    row_num = 1
    for row in initial_rows:
        extra_row_for_work = '[작업사항]' in row
        if extra_row_for_work:
            sheet.insert_rows(row_num-1)
            sheet.cell(row=row_num-1, column=1, value='띄어쓰기')
            row_num += 1  # Adjusting row_num because of the new row
        row_num += 1  # Proceed to the next original row

    # Save the modified workbook
    workbook.save(new_file_path)

# 형태 변경된 worklist 마지막 전처리
def drop_first_row(file_path, new_file_path):
    # Open the final Excel file and drop the first row
    df = pd.read_excel(file_path)
    df = df.drop(index=[0])
    with pd.ExcelWriter(new_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
        
# worklist 그루핑 함수        
def group_and_save_data(uploaded_file, file_type):
    # Load data depending on file type
    if file_type == 'csv':
        df = pd.read_csv(uploaded_file)
    elif file_type == 'xlsx':
        df = pd.read_excel(uploaded_file, engine='openpyxl')
    else:
        return "Unsupported file format."

    # Function to check whether a line has a ‘space’
    def has_spacing(row):
        return any('띄어쓰기' in str(cell) for cell in row)

    # Initialize variables for group creation
    groups = []
    start = 0

    # Traverse rows of the data frame
    for idx, row in df.iterrows():
        if has_spacing(row):
            group = df.iloc[start:idx+1]  # Slice up to the current row and create a group.
            groups.append(group)
            start = idx + 1  # Set the starting point for the next group

    # Add last section (group)
    last_section = df.iloc[start:]
    groups.append(last_section)

    # Use tempfile to create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        with pd.ExcelWriter(tmp.name, engine='openpyxl') as writer:
            for i, group in enumerate(groups):
                group.to_excel(writer, index=False, sheet_name=f'Group_{i+1}')
        
        # Return the path of the temporary file
        return tmp.name
    
def categorize_sheet(sheet, categories):

    cell_value = sheet.iloc[0, 2]  # This is equivalent to 'C1' in 1-based index (Excel style)

    for category, identifiers in categories.items():
        if cell_value in identifiers:
            return category

    return None  # If no category matches

def save_to_excel(groups, file_name):

    with pd.ExcelWriter(file_name) as writer:
        for name, group in groups.items():
            for idx, df in enumerate(group, 1):
                df.to_excel(writer, sheet_name=f'{name}_{idx}', index=False)
                
# 장치 속성 별 그루핑 함수              
def main_grouping(grouped_file_path):
    # Load the previously created Excel file with all the groups
    xls = pd.ExcelFile(grouped_file_path)
    
    # Categories and their respective identifiers
    categories = {
        'CDV': ['C', 'D', 'F'],
        'HEX': ['HEX'],
        'AFC': ['AFC'],
    }

    # We will store the sheets in these groups based on the category they fall into
    categorized_sheets = {key: [] for key in categories.keys()}

    # Iterate through all sheets in the Excel file
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)

        # Determine the category of the current sheet
        category = categorize_sheet(df, categories)

        # If the sheet belongs to a category, add it to the corresponding group
        if category:
            categorized_sheets[category].append(df)
    
    # Prepare the files for download but do not create download links yet
    file_paths = {}
    for category, sheets in categorized_sheets.items():
        if sheets:  # Check if there are sheets in the category
            file_name = f'{category}_grouped.xlsx'
            save_to_excel({category: sheets}, file_name)
            file_paths[category] = file_name

    return file_paths

# Checklist_Temp 폴더 자동 삭제 함수
def clear_directory(directory):
    """
    Clears all files and folders in the specified directory.
    """
    for item in os.listdir(directory):
        item_path = os.path.join(directory, item)
        if os.path.isfile(item_path):
            os.remove(item_path)
        elif os.path.isdir(item_path):
            shutil.rmtree(item_path)

# cdv checklist 자동생성 함수
def cdv_checklist(uploaded_file, filename1):
    # 파일 로드
    grouped_data = pd.read_excel(uploaded_file, sheet_name=None)

    output_folder = "Checklist_Temp"
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)
    
    # 각 그룹(시트)에 대해 반복
    for sheet_name, sheet_data in grouped_data.items():
        # 새로운 체크리스트 워크북을 만듭니다 (위에서 제공한 코드를 계속 사용합니다).
        # ... [워크북 생성 및 서식 지정 코드] ...

        # Create a new Excel workbook and get the active worksheet
        wb = openpyxl.Workbook()
        ws = wb.active

        # Define the number of rows and columns
        num_rows = 66
        num_columns = 8  # Columns A to H

        # Populate the table with placeholder data
        for row in range(1, num_rows + 1):
            for col in range(1, num_columns + 1):
                cell = ws[get_column_letter(col) + str(row)]
                cell.value = None

        # 열 너비 설정
        column_widths = {'A': 9.9,'B': 24.3,'C': 7.9,'D': 29.9,'E': 9.9,'F': 17.5,'G': 7.9,'H': 29.9}

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # 행 높이 설정
        row_heights = {1: 15, 6: 15, 2: 37.8}

        # Apply the specified heights
        for row_num, height in row_heights.items():
            ws.row_dimensions[row_num].height = height

        # Set the height for all other rows
        for row_num in range(1, num_rows + 1):
            if row_num not in row_heights:
                ws.row_dimensions[row_num].height = 30

        # 2행 병합
        ws.merge_cells('A2:H2')

        # 2행 문자 삽입
        cell_2 = ws['A2']
        cell_2.value = "압력용기 개방검사 검사표"
        cell_2.font = Font(bold=True, underline='single', size=25)
        cell_2.alignment = Alignment(horizontal='center', vertical='center')

        # 3~5행 병합
        ws.merge_cells('C3:E3')
        ws.merge_cells('C4:E4')
        ws.merge_cells('C5:E5')

        # 3행 문자 삽입
        cell_3_a = ws['A3']
        cell_3_a.value = '장치번호'
        cell_3_a.font = Font(bold=True, size=14)
        cell_3_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_3_f = ws['F3']
        cell_3_f.value = '장치등급'
        cell_3_f.font = Font(bold=True, size=14)
        cell_3_f.alignment = Alignment(horizontal='left', vertical='center')

        # 4행 문자 삽입
        cell_4_a = ws['A4']
        cell_4_a.value = '검사일 :'
        cell_4_a.font = Font(bold=True, size=14)
        cell_4_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_4_f = ws['F4']
        cell_4_f.value = '검사구분'
        cell_4_f.font = Font(bold=True, size=14)
        cell_4_f.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 문자 삽입
        cell_5_a = ws['A5']
        cell_5_a.value = '검사원 :'
        cell_5_a.font = Font(bold=True, size=14)
        cell_5_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_c = ws['C5']
        cell_5_c.value = '(서명)'
        cell_5_c.font = Font(bold=True, size=11)
        cell_5_c.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_f = ws['F5']
        cell_5_f.value = '엔지니어 :'
        cell_5_f.font = Font(bold=True, size=14)
        cell_5_f.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_h = ws['H5']
        cell_5_h.value = '(서명)'
        cell_5_h.font = Font(bold=True, size=11)
        cell_5_h.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 아래 굵은 테두리 표시
        bold_bottom_border = Border(bottom=Side(style='thick'))

        for col in range(1, 9):
            ws.cell(row=5, column=col).border = bold_bottom_border

        # 7행 병합 및 문자 삽입
        ws.merge_cells('A7:H7')
        cell_7_a = ws['A7']
        cell_7_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_7_a.value = '1. TA Worklist'
        cell_7_a.font = Font(bold=True, size=12)
        cell_7_a.alignment = Alignment(horizontal='left', vertical='center')

        # 8행 병합 및 문자 삽입
        ws.merge_cells('B8:H8')
        cell_8_a = ws['A8']
        cell_8_a.value = '번호'
        cell_8_a.font = Font(bold=True, size=12)
        cell_8_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_8_b = ws['B8']
        cell_8_b.value = '작업 사항'
        cell_8_b.font = Font(bold=True, size=12)
        cell_8_b.alignment = Alignment(horizontal='center', vertical='center')

        # 9~18행 병합 및 문자 삽입
        for i in range(9, 19):
            ws.merge_cells(f'B{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 8)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 19행 병합 및 문자 삽입
        ws.merge_cells('A19:H19')
        cell_19_a = ws['A19']
        cell_19_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_19_a.value = '2. Damage Mechanism :'
        cell_19_a.font = Font(bold=True, size=12)
        cell_19_a.alignment = Alignment(horizontal='left', vertical='center')

        # 20행 병합 및 문자 삽입
        ws.merge_cells('A20:H20')
        cell_20_a = ws['A20']
        cell_20_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_20_a.value = '3. 중요 검사 항목'
        cell_20_a.font = Font(bold=True, size=12)
        cell_20_a.alignment = Alignment(horizontal='left', vertical='center')

        # 21행 병합 및 문자 삽입
        ws.merge_cells('B21:H21')
        cell_21_a = ws['A21']
        cell_21_a.value = '번호'
        cell_21_a.font = Font(bold=True, size=12)
        cell_21_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_b = ws['B21']
        cell_21_b.value = '검사 항목'
        cell_21_b.font = Font(bold=True, size=12)
        cell_21_b.alignment = Alignment(horizontal='center', vertical='center')

        # 22행~31행 병합 및 문자 삽입
        for i in range(22, 32):
            ws.merge_cells(f'B{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 21)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 32행 병합 및 문자 삽입
        ws.merge_cells('A32:H32')
        cell_32_a = ws['A32']
        cell_32_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_32_a.value = '4. 주요 검사 이력 (전기 TA, 운전 중 이력)'
        cell_32_a.font = Font(bold=True, size=12)
        cell_32_a.alignment = Alignment(horizontal='left', vertical='center')

        # 33행 병합 및 문자 삽입
        ws.merge_cells('B33:H33')
        cell_33_a = ws['A33']
        cell_33_a.value = '번호'
        cell_33_a.font = Font(bold=True, size=12)
        cell_33_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_33_b = ws['B33']
        cell_33_b.value = '검사 이력'
        cell_33_b.font = Font(bold=True, size=12)
        cell_33_b.alignment = Alignment(horizontal='center', vertical='center')

        # 34행~43행 병합 및 문자 삽입
        for i in range(34, 44):
            ws.merge_cells(f'B{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 33)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 44행 병합 및 문자 삽입
        ws.merge_cells('A44:H44')
        cell_44_a = ws['A44']
        cell_44_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_44_a.value = '5. 장치 Type별 일반 검사 항목 (양호, 불량, N/A)'
        cell_44_a.font = Font(bold=True, size=12)
        cell_44_a.alignment = Alignment(horizontal='left', vertical='center')

        # 45행 병합 및 문자 삽입
        cell_45_a = ws['A45']
        cell_45_a.value = '구분'
        cell_45_a.font = Font(bold=True, size=12)
        cell_45_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_b = ws['B45']
        cell_45_b.value = '검사항목'
        cell_45_b.font = Font(bold=True, size=12)
        cell_45_b.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('C45:D45')
        cell_45_c = ws['C45']
        cell_45_c.value = '검사 결과'
        cell_45_c.font = Font(bold=True, size=12)
        cell_45_c.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_e = ws['E45']
        cell_45_e.value = '구분'
        cell_45_e.font = Font(bold=True, size=12)
        cell_45_e.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_f = ws['F45']
        cell_45_f.value = '검사항목'
        cell_45_f.font = Font(bold=True, size=12)
        cell_45_f.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('G45:H45')
        cell_45_g = ws['G45']
        cell_45_g.value = '검사 결과'
        cell_45_g.font = Font(bold=True, size=12)
        cell_45_g.alignment = Alignment(horizontal='center', vertical='center')

        # 46행 문자 삽입
        cell_46_b = ws['B46']
        cell_46_b.value = '청소 전 초기상태'
        cell_46_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_46_f= ws['F46']
        cell_46_f.value = 'Tray Support'
        cell_46_f.alignment = Alignment(horizontal='center', vertical='center')

        # 47행 병합 및 문자 삽입
        cell_47_b = ws['B47']
        cell_47_b.value = 'Sludge, Scale 부착정도'
        cell_47_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_47_f= ws['F47']
        cell_47_f.value = 'Down Comer'
        cell_47_f.alignment = Alignment(horizontal='center', vertical='center')

        # 문자 삽입
        cell_48_b = ws['B48']
        cell_48_b.value = 'Shell'
        cell_48_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_48_f= ws['F48']
        cell_48_f.value = 'Draw Off'
        cell_48_f.alignment = Alignment(horizontal='center', vertical='center')

        # 49 행 문자 삽입
        cell_49_b = ws['B49']
        cell_49_b.value = 'Head'
        cell_49_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_49_f= ws['F49']
        cell_49_f.value = 'Entry Horn'
        cell_49_f.alignment = Alignment(horizontal='center', vertical='center')

        # 50행 문자 삽입
        cell_50_b = ws['B50']
        cell_50_b.value = 'Boot'
        cell_50_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_50_f= ws['F50']
        cell_50_f.value = 'Packing'
        cell_50_f.alignment = Alignment(horizontal='center', vertical='center')

        # 51행 문자 삽입
        cell_51_b = ws['B51']
        cell_51_b.value = 'Nozzle & Flange'
        cell_51_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_51_f= ws['F51']
        cell_51_f.value = 'Electrode'
        cell_51_f.alignment = Alignment(horizontal='center', vertical='center')

        # 52행 문자 삽입
        cell_52_b = ws['B52']
        cell_52_b.value = 'Lining/Paint 상태'
        cell_52_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_52_f= ws['F52']
        cell_52_f.value = 'Grinding'
        cell_52_f.alignment = Alignment(horizontal='center', vertical='center')

        # 53행 문자 삽입
        cell_53_b = ws['B53']
        cell_53_b.value = 'Main 용접심'
        cell_53_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_53_f= ws['F53']
        cell_53_f.value = '용접'
        cell_53_f.alignment = Alignment(horizontal='center', vertical='center')

        # 54행 문자 삽입
        cell_54_b = ws['B54']
        cell_54_b.value = '부착물 용접심'
        cell_54_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_54_f= ws['F54']
        cell_54_f.value = '열처리'
        cell_54_f.alignment = Alignment(horizontal='center', vertical='center')

        # 55행 문자 삽입
        cell_55_b = ws['B55']
        cell_55_b.value = 'Shell to Nozzle 용접심'
        cell_55_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_55_f= ws['F55']
        cell_55_f.value = '보수부위 NDT'
        cell_55_f.alignment = Alignment(horizontal='center', vertical='center')

        # 56행 문자 삽입
        cell_56_b = ws['B56']
        cell_56_b.value = '노즐 Sleeve'
        cell_56_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_56_f= ws['F56']
        cell_56_f.value = '압력시험'
        cell_56_f.alignment = Alignment(horizontal='center', vertical='center')

        # 57행 문자 삽입
        cell_57_b = ws['B57']
        cell_57_b.value = 'Distributor / Collector'
        cell_57_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_57_f= ws['F57']
        cell_57_f.value = 'Painting'
        cell_57_f.alignment = Alignment(horizontal='center', vertical='center')

        # 58행 문자 삽입
        cell_58_b = ws['B58']
        cell_58_b.value = 'Internal Piping'
        cell_58_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_58_f= ws['F58']
        cell_58_f.value = 'Tray M/W'
        cell_58_f.alignment = Alignment(horizontal='center', vertical='center')

        # 59행 문자 삽입
        cell_59_b = ws['B59']
        cell_59_b.value = 'Johnson Screen'
        cell_59_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_59_f= ws['F59']
        cell_59_f.value = 'Distributor Pipe'
        cell_59_f.alignment = Alignment(horizontal='center', vertical='center')

        # 60행 문자 삽입
        cell_60_b = ws['B60']
        cell_60_b.value = 'Wire Mesh'
        cell_60_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_60_f= ws['F60']
        cell_60_f.value = 'Shell/Head'
        cell_60_f.alignment = Alignment(horizontal='center', vertical='center')

        # 61행 문자 삽입
        cell_61_b = ws['B61']
        cell_61_b.value = 'Grid'
        cell_61_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_61_f= ws['F61']
        cell_61_f.value = 'Nozzle'
        cell_61_f.alignment = Alignment(horizontal='center', vertical='center')

        # 62행 문자 삽입
        cell_62_b = ws['B62']
        cell_62_b.value = 'Screen / Grid Support'
        cell_62_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_62_f= ws['F62']
        cell_62_f.value = '보온재'
        cell_62_f.alignment = Alignment(horizontal='center', vertical='center')

        # 63행 문자 삽입
        cell_63_b = ws['B63']
        cell_63_b.value = 'Vortex Breaker'
        cell_63_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_63_f= ws['F63']
        cell_63_f.value = 'Nozzle & Flange'
        cell_63_f.alignment = Alignment(horizontal='center', vertical='center')

        # 64행 문자 삽입
        cell_64_b = ws['B64']
        cell_64_b.value = 'Tray'
        cell_64_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_64_f= ws['F64']
        cell_64_f.value = 'Fittings'
        cell_64_f.alignment = Alignment(horizontal='center', vertical='center')

        # 65행 문자 삽입
        cell_65_b = ws['B65']
        cell_65_b.value = 'Tray Cap'
        cell_65_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_65_f= ws['F65']
        cell_65_f.value = 'Supports'
        cell_65_f.alignment = Alignment(horizontal='center', vertical='center')

        # 66행 문자 삽입
        cell_66_b = ws['B66']
        cell_66_b.value = 'Tray Bolt/Nut/Clamp'
        cell_66_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_66_f= ws['F66']
        cell_66_f.value = 'Fire Proofing'
        cell_66_f.alignment = Alignment(horizontal='center', vertical='center')

        # 세로 병합
        ws.merge_cells('A46:A52')
        cell_46_a = ws['A46']
        cell_46_a.value = '내부\n육안검사'
        cell_46_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E46:E51')
        cell_46_e = ws['E46']
        cell_46_e.value = 'Internals\n상태검사'
        cell_46_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A53:A57')
        cell_53_a = ws['A53']
        cell_53_a.value = '내부 NDT'
        cell_53_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E52:E57')
        cell_52_e = ws['E52']
        cell_52_e.value = '추가\n보수작업'
        cell_52_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A58:A66')
        cell_58_a = ws['A58']
        cell_58_a.value = 'Internals\n상태검사'
        cell_58_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E58:E59')
        cell_58_e = ws['E58']
        cell_58_e.value = 'Internal\n조립'
        cell_58_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E60:E61')
        cell_60_e = ws['E60']
        cell_60_e.value = '두께측정'
        cell_60_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E62:E66')
        cell_62_e = ws['E62']
        cell_62_e.value = '외부\n육안 검사'
        cell_62_e.alignment = Alignment(horizontal='center', vertical='center')

        # 테두리 적용
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))

        # Apply the border to all cells in the range
        for row in ws.iter_rows(min_row=1, max_row=num_rows, min_col=1, max_col=num_columns):
            for cell in row:
                cell.border = thin_border

    # =============================================================================================================
        # 장치 번호 입력

        device_number = sheet_data.iloc[0, 0]  # 장치 번호 추정 위치
        ws['B3'].value = device_number  # 장치 번호 셀 할당
        ws['B3'].alignment = Alignment(horizontal='center', vertical='center')

    # =============================================================================================================    

        # 작업 사항 및 검사 사항 사이의 데이터를 처리하기 위한 로직
        worklist_start = None
        content_end = None  # 작업 사항과 검사 사항 사이의 내용을 저장할 위치
        separate_row_found = False  # '띄어쓰기' 행이 발견되었는지 확인

        for i, row in sheet_data.iterrows():
            # '[작업사항]'이 발견되면 시작 위치를 설정합니다.
            if '[작업사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                worklist_start = i

            # '[검사사항]'이 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '[검사사항]' in str(row[0]) and worklist_start is not None:
                content_end = i
                break  # '[검사사항]'이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

            # 띄어쓰기가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif row[0] == '띄어쓰기' and worklist_start is not None:  # row[0]이 빈 문자열인 경우를 확인합니다.
                content_end = i
                empty_row_found = True  # 빈 행이 발견되었음을 표시합니다.
                break  # 빈 행이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 작업 사항과 검사 사항 사이에 데이터가 있고, '[검사사항]' 또는 빈 행이 발견된 경우 데이터 삽입을 시작합니다.
        if worklist_start is not None and content_end is not None:
            worklist_data = sheet_data.iloc[worklist_start + 1:content_end]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(worklist_data[first_column_name], start=9):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

    # =============================================================================================================  
        # 검사 사항과 '띄어쓰기' 사이의 데이터를 처리하기 위한 로직
        inpection_start = None
        content_end1 = None  # 검사 사항과 띄어쓰기 사이의 내용을 저장할 위치

        for i, row in sheet_data.iterrows():
            # '[검사사항]'이 발견되면 시작 위치를 설정합니다.
            if '[검사사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                inpection_start = i

            # '띄어쓰기'가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '띄어쓰기' in str(row[0]) and inpection_start is not None:
                content_end1 = i
                break  # '띄어쓰기'가 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 검사 사항과 '띄어쓰기' 사이에 데이터 삽입을 시작합니다.
        if inpection_start is not None and content_end1 is not None:
            inpection_data = sheet_data.iloc[inpection_start + 1:content_end1]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name1 = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(inpection_data[first_column_name1], start=22):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기


        # 파일을 지정된 폴더에 저장합니다.

        # 각 파일에 대한 개별 폴더 생성
        filename1 = filename1
        filename2 = device_number.replace("/", "&").strip()
        filename3 = 'Photo'
        filename4 = '도면'

        individual_folder1 = os.path.join(output_folder, filename1)
        if not os.path.exists(individual_folder1):
            os.mkdir(individual_folder1)

        individual_folder2 = os.path.join(individual_folder1, filename2)
        if not os.path.exists(individual_folder2):
            os.mkdir(individual_folder2)

        individual_folder3 = os.path.join(individual_folder2, filename3)
        if not os.path.exists(individual_folder3):
            os.mkdir(individual_folder3)

        individual_folder4 = os.path.join(individual_folder2, filename4)
        if not os.path.exists(individual_folder4):
            os.mkdir(individual_folder4)

        wb.save(os.path.join(individual_folder2, f"Checklist_{filename2}.xlsx"))

    # 생성된 모든 파일들을 하나의 압축 파일로 만듭니다.
    shutil.make_archive(output_folder, 'zip', output_folder)
    return f"{output_folder}.zip"

# hex checklist 자동생성 함수
def hex_checklist(uploaded_file, filename1):
    # 파일 로드
    grouped_data = pd.read_excel(uploaded_file, sheet_name=None)

    output_folder = "Checklist_Temp"
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)
    
    # 각 그룹(시트)에 대해 반복
    for sheet_name, sheet_data in grouped_data.items():
        # 새로운 체크리스트 워크북을 만듭니다 (위에서 제공한 코드를 계속 사용합니다).
        # ... [워크북 생성 및 서식 지정 코드] ...

        # Create a new Excel workbook and get the active worksheet
        wb = openpyxl.Workbook()
        ws = wb.active

        # Define the number of rows and columns
        num_rows = 45
        num_columns = 8  # Columns A to H

        # Populate the table with placeholder data
        for row in range(1, num_rows + 1):
            for col in range(1, num_columns + 1):
                cell = ws[get_column_letter(col) + str(row)]
                cell.value = None

        # 열 너비 설정
        column_widths = {'A': 14,'B': 35.5,'C': 7.9,'D': 29,'E': 12,'F': 28,'G': 21,'H': 29}

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # 행 높이 설정
        row_heights = {1: 15, 6: 15, 2: 37.8}

        # Apply the specified heights
        for row_num, height in row_heights.items():
            ws.row_dimensions[row_num].height = height

        # Set the height for all other rows
        for row_num in range(1, num_rows + 1):
            if row_num not in row_heights:
                ws.row_dimensions[row_num].height = 30

        # 2행 병합
        ws.merge_cells('A2:H2')

        # 2행 문자 삽입
        cell_2 = ws['A2']
        cell_2.value = "열교환기 개방검사 검사표"
        cell_2.font = Font(bold=True, underline='single', size=25)
        cell_2.alignment = Alignment(horizontal='center', vertical='center')

        # 3~5행 병합
        ws.merge_cells('C3:E3')
        ws.merge_cells('C4:E4')
        ws.merge_cells('C5:E5')

        # 3행 문자 삽입
        cell_3_a = ws['A3']
        cell_3_a.value = '장치번호'
        cell_3_a.font = Font(bold=True, size=14)
        cell_3_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_3_f = ws['F3']
        cell_3_f.value = '장치등급'
        cell_3_f.font = Font(bold=True, size=14)
        cell_3_f.alignment = Alignment(horizontal='left', vertical='center')

        # 4행 문자 삽입
        cell_4_a = ws['A4']
        cell_4_a.value = '검사일 :'
        cell_4_a.font = Font(bold=True, size=14)
        cell_4_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_4_f = ws['F4']
        cell_4_f.value = '검사구분'
        cell_4_f.font = Font(bold=True, size=14)
        cell_4_f.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 문자 삽입
        cell_5_a = ws['A5']
        cell_5_a.value = '검사원 :'
        cell_5_a.font = Font(bold=True, size=14)
        cell_5_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_c = ws['C5']
        cell_5_c.value = '(서명)'
        cell_5_c.font = Font(bold=True, size=11)
        cell_5_c.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_f = ws['F5']
        cell_5_f.value = '엔지니어 :'
        cell_5_f.font = Font(bold=True, size=14)
        cell_5_f.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_h = ws['H5']
        cell_5_h.value = '(서명)'
        cell_5_h.font = Font(bold=True, size=11)
        cell_5_h.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 아래 굵은 테두리 표시
        bold_bottom_border = Border(bottom=Side(style='thick'))

        for col in range(1, 9):
            ws.cell(row=5, column=col).border = bold_bottom_border

        # 7행 병합 및 문자 삽입
        ws.merge_cells('A7:H7')
        cell_7_a = ws['A7']
        cell_7_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_7_a.value = '1. TA Worklist'
        cell_7_a.font = Font(bold=True, size=12)
        cell_7_a.alignment = Alignment(horizontal='left', vertical='center')

        # 8행 병합 및 문자 삽입
        ws.merge_cells('B8:E8')
        cell_8_a = ws['A8']
        cell_8_a.value = '번호'
        cell_8_a.font = Font(bold=True, size=12)
        cell_8_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('F8:H8')
        cell_8_f = ws['F8']
        cell_8_f.value = '주요 검사 사항'
        cell_8_f.font = Font(bold=True, size=12)
        cell_8_f.alignment = Alignment(horizontal='center', vertical='center')

        cell_8_b = ws['B8']
        cell_8_b.value = '작업 내용'
        cell_8_b.font = Font(bold=True, size=12)
        cell_8_b.alignment = Alignment(horizontal='center', vertical='center')

        # 9~18행 병합 및 문자 삽입
        for i in range(9, 19):
            ws.merge_cells(f'B{i}:E{i}')
            ws.merge_cells(f'F{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 8)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')      

        # 19행 병합 및 문자 삽입
        ws.merge_cells('A19:H19')
        cell_19_a = ws['A19']
        cell_19_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_19_a.value = '2. Damage Mechanism :'
        cell_19_a.font = Font(bold=True, size=12)
        cell_19_a.alignment = Alignment(horizontal='left', vertical='center')

        # 20행 병합 및 문자 삽입
        ws.merge_cells('A20:H20')
        cell_20_a = ws['A20']
        cell_20_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_20_a.value = '3. 장치 Type별 일반 검사 항목 (양호, 불량, N/A)'
        cell_20_a.font = Font(bold=True, size=12)
        cell_20_a.alignment = Alignment(horizontal='left', vertical='center')

        # 21행 병합 및 문자 삽입
        cell_21_a = ws['A21']
        cell_21_a.value = '구분'
        cell_21_a.font = Font(bold=True, size=12)
        cell_21_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_b = ws['B21']
        cell_21_b.value = '검사항목'
        cell_21_b.font = Font(bold=True, size=12)
        cell_21_b.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('C21:D21')
        cell_21_c = ws['C21']
        cell_21_c.value = '검사 결과'
        cell_21_c.font = Font(bold=True, size=12)
        cell_21_c.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_e = ws['E21']
        cell_21_e.value = '구분'
        cell_21_e.font = Font(bold=True, size=12)
        cell_21_e.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_f = ws['F21']
        cell_21_f.value = '검사항목'
        cell_21_f.font = Font(bold=True, size=12)
        cell_21_f.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('G21:H21')
        cell_21_g = ws['G21']
        cell_21_g.value = '검사 결과'
        cell_21_g.font = Font(bold=True, size=12)
        cell_21_g.alignment = Alignment(horizontal='center', vertical='center')

        # 22행 문자 삽입
        cell_22_b = ws['B22']
        cell_22_b.value = '청소 전 초기상태'
        cell_22_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_22_f= ws['F22']
        cell_22_f.value = 'Shell & Head(Header Box)'
        cell_22_f.alignment = Alignment(horizontal='center', vertical='center')

        # 23행 병합 및 문자 삽입
        cell_23_b = ws['B23']
        cell_23_b.value = 'Sludge, Scale 부착정도'
        cell_23_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_23_f= ws['F23']
        cell_23_f.value = 'Channel & Cover'
        cell_23_f.alignment = Alignment(horizontal='center', vertical='center')

        # 24행 병합 및 문자 삽입
        cell_24_b = ws['B24']
        cell_24_b.value = 'Shell & Head'
        cell_24_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_24_f= ws['F24']
        cell_24_f.value = 'Nozzle & Flange'
        cell_24_f.alignment = Alignment(horizontal='center', vertical='center')

        # 25행 문자 삽입
        cell_25_b = ws['B25']
        cell_25_b.value = 'Channel & Cover'
        cell_25_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_25_f= ws['F25']
        cell_25_f.value = 'IRIS/ECT (본)'
        cell_25_f.alignment = Alignment(horizontal='center', vertical='center')

        # 26행 문자 삽입
        cell_26_b = ws['B26']
        cell_26_b.value = 'Nozzle & Flange'
        cell_26_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_26_f= ws['F26']
        cell_26_f.value = 'Shell/Channel Main 용접심'
        cell_26_f.alignment = Alignment(horizontal='center', vertical='center')

        # 27행 문자 삽입
        cell_27_b = ws['B27']
        cell_27_b.value = 'Floating Head'
        cell_27_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_27_f = ws['F27']
        cell_27_f.value = 'Shell/Channel 부착물 용접심'
        cell_27_f.alignment = Alignment(horizontal='center', vertical='center')

        # 28행 문자 삽입
        cell_28_b = ws['B28']
        cell_28_b.value = 'Header Box'
        cell_28_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_28_f = ws['F28']
        cell_28_f.value = 'Shell/Channel to Nozzle 용접심'
        cell_28_f.alignment = Alignment(horizontal='center', vertical='center')

        # 29행 문자 삽입
        cell_29_b = ws['B29']
        cell_29_b.value = 'Header Box 나사산'
        cell_29_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_29_f= ws['F29']
        cell_29_f.value = 'Nozzle C & B Joint'
        cell_29_f.alignment = Alignment(horizontal='center', vertical='center')

        # 30행 문자 삽입
        cell_30_b = ws['B30']
        cell_30_b.value = 'Plug Bolt 나사산'
        cell_30_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_30_f= ws['F30']
        cell_30_f.value = 'U-Bend RT'
        cell_30_f.alignment = Alignment(horizontal='center', vertical='center')

        # 31행 문자 삽입
        cell_31_b = ws['B31']
        cell_31_b.value = 'Plug Gasket Face'
        cell_31_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_31_f= ws['F31']
        cell_31_f.value = '상부 1열 RT(Fixed Type)'
        cell_31_f.alignment = Alignment(horizontal='center', vertical='center')

        # 32행 문자 삽입
        cell_32_b = ws['B32']
        cell_32_b.value = 'Coating/Lining 상태'
        cell_32_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_32_f= ws['F32']
        cell_32_f.value = 'Strength Welding부 PT'
        cell_32_f.alignment = Alignment(horizontal='center', vertical='center')

        # 33행 문자 삽입
        cell_33_b = ws['B33']
        cell_33_b.value = 'Tube 외부'
        cell_33_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_33_f= ws['F33']
        cell_33_f.value = 'Header Box 용접심'
        cell_33_f.alignment = Alignment(horizontal='center', vertical='center')

        # 34행 문자 삽입
        cell_34_b = ws['B34']
        cell_34_b.value = 'Tube 내부'
        cell_34_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_34_f= ws['F34']
        cell_34_f.value = 'Partition/Stiffener Plate 용접부'
        cell_34_f.alignment = Alignment(horizontal='center', vertical='center')

        # 35행 문자 삽입
        cell_35_b = ws['B35']
        cell_35_b.value = 'Tube Sheet'
        cell_35_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_35_f= ws['F35']
        cell_35_f.value = '용접'
        cell_35_f.alignment = Alignment(horizontal='center', vertical='center')

        # 36행 문자 삽입
        cell_36_b = ws['B36']
        cell_36_b.value = 'Baffle, Tie-Rods'
        cell_36_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_36_f= ws['F36']
        cell_36_f.value = '열처리'
        cell_36_f.alignment = Alignment(horizontal='center', vertical='center')

        # 37행 문자 삽입
        cell_37_b = ws['B37']
        cell_37_b.value = 'Sealing Strip'
        cell_37_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_37_f = ws['F37']
        cell_37_f.value = 'Grinding'
        cell_37_f.alignment = Alignment(horizontal='center', vertical='center')

        # 38행 문자 삽입
        cell_38_b = ws['B38']
        cell_38_b.value = 'Impingement Plate'
        cell_38_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_38_f = ws['F38']
        cell_38_f.value = '보수부위 NDT'
        cell_38_f.alignment = Alignment(horizontal='center', vertical='center')

        # 39행 문자 삽입
        cell_39_b = ws['B39']
        cell_39_b.value = '보온재'
        cell_39_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_39_f = ws['F39']
        cell_39_f.value = '압력시험'
        cell_39_f.alignment = Alignment(horizontal='center', vertical='center')

        # 40행 문자 삽입
        cell_40_b = ws['B40']
        cell_40_b.value = 'Nozzle & Flange'
        cell_40_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_40_f = ws['F40']
        cell_40_f.value = 'Painting'
        cell_40_f.alignment = Alignment(horizontal='center', vertical='center')

        # 41행 문자 삽입
        cell_41_b = ws['B41']
        cell_41_b.value = 'Fittings'
        cell_41_b.alignment = Alignment(horizontal='center', vertical='center')

        # 42행 문자 삽입
        cell_42_b = ws['B42']
        cell_42_b.value = 'Supports'
        cell_42_b.alignment = Alignment(horizontal='center', vertical='center')

        # 43행 문자 삽입
        cell_43_b = ws['B43']
        cell_43_b.value = 'Fire Proofing'
        cell_43_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_43_f = ws['F43']
        cell_43_f.value = '1차'
        cell_43_f.alignment = Alignment(horizontal='center', vertical='center')

        # 44행 문자 삽입
        cell_44_b = ws['B44']
        cell_44_b.value = 'Paint 상태'
        cell_44_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_44_f = ws['F44']
        cell_44_f.value = '2차'
        cell_44_f.alignment = Alignment(horizontal='center', vertical='center')

        # 45행 문자 삽입
        cell_45_b = ws['B45']
        cell_45_b.value = 'Saddle Sliding Side\nDe-Bolting 유무 점검'
        cell_45_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_f = ws['F45']
        cell_45_f.value = '3차'
        cell_45_f.alignment = Alignment(horizontal='center', vertical='center')

        # 세로 병합
        ws.merge_cells('A22:A32')
        cell_22_a = ws['A22']
        cell_22_a.value = '내부\n육안검사'
        cell_22_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E22:E25')
        cell_22_e = ws['E22']
        cell_22_e.value = '두께측정'
        cell_22_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A33:A38')
        cell_33_a = ws['A33']
        cell_33_a.value = 'Bundle\n육안검사'
        cell_33_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E26:E34')
        cell_26_e = ws['E26']
        cell_26_e.value = '내부\nNDT'
        cell_26_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A39:A45')
        cell_39_a = ws['A39']
        cell_39_a.value = '외부 육안\n검사'
        cell_39_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E35:E42')
        cell_35_e = ws['E35']
        cell_35_e.value = '추가\n보수작업'
        cell_35_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E43:E45')
        cell_43_e = ws['E43']
        cell_43_e.value = '수압시험'
        cell_43_e.alignment = Alignment(horizontal='center', vertical='center')

        # 테두리 적용
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))

        # Apply the border to all cells in the range
        for row in ws.iter_rows(min_row=1, max_row=num_rows, min_col=1, max_col=num_columns):
            for cell in row:
                cell.border = thin_border
      ########################################################################### 여기부터 수정 필요           
    # =============================================================================================================

        # 장치 번호 입력
        device_number = sheet_data.iloc[0, 0]  # 장치 번호 추정 위치
        ws['B3'].value = device_number  # 장치 번호 셀 할당
        ws['B3'].alignment = Alignment(horizontal='center', vertical='center')

    # =============================================================================================================    

        # 작업 사항 및 검사 사항 사이의 데이터를 처리하기 위한 로직
        worklist_start = None
        content_end = None  # 작업 사항과 검사 사항 사이의 내용을 저장할 위치
        separate_row_found = False  # '띄어쓰기' 행이 발견되었는지 확인

        for i, row in sheet_data.iterrows():
            # '[작업사항]'이 발견되면 시작 위치를 설정합니다.
            if '[작업사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                worklist_start = i

            # '[검사사항]'이 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '[검사사항]' in str(row[0]) and worklist_start is not None:
                content_end = i
                break  # '[검사사항]'이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

            # 띄어쓰기가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif row[0] == '띄어쓰기' and worklist_start is not None:  # row[0]이 빈 문자열인 경우를 확인합니다.
                content_end = i
                empty_row_found = True  # 빈 행이 발견되었음을 표시합니다.
                break  # 빈 행이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 작업 사항과 검사 사항 사이에 데이터가 있고, '[검사사항]' 또는 빈 행이 발견된 경우 데이터 삽입을 시작합니다.
        if worklist_start is not None and content_end is not None:
            worklist_data = sheet_data.iloc[worklist_start + 1:content_end]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(worklist_data[first_column_name], start=9):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

    # =============================================================================================================  
        # 검사 사항과 '띄어쓰기' 사이의 데이터를 처리하기 위한 로직
        inpection_start = None
        content_end1 = None  # 검사 사항과 띄어쓰기 사이의 내용을 저장할 위치

        for i, row in sheet_data.iterrows():
            # '[검사사항]'이 발견되면 시작 위치를 설정합니다.
            if '[검사사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                inpection_start = i

            # '띄어쓰기'가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '띄어쓰기' in str(row[0]) and inpection_start is not None:
                content_end1 = i
                break  # '띄어쓰기'가 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 검사 사항과 '띄어쓰기' 사이에 데이터 삽입을 시작합니다.
        if inpection_start is not None and content_end1 is not None:
            inpection_data = sheet_data.iloc[inpection_start + 1:content_end1]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name1 = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # F9셀부터 데이터 삽입
            for i, item in enumerate(inpection_data[first_column_name1], start=9):
                cell = ws.cell(row=i, column=6)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기


        # 파일을 지정된 폴더에 저장합니다.

        # 각 파일에 대한 개별 폴더 생성
        filename1 = filename1
        filename2 = device_number.replace("/", "&").strip()
        filename3 = 'Photo'
        filename4 = '도면'

        individual_folder1 = os.path.join(output_folder, filename1)
        if not os.path.exists(individual_folder1):
            os.mkdir(individual_folder1)

        individual_folder2 = os.path.join(individual_folder1, filename2)
        if not os.path.exists(individual_folder2):
            os.mkdir(individual_folder2)

        individual_folder3 = os.path.join(individual_folder2, filename3)
        if not os.path.exists(individual_folder3):
            os.mkdir(individual_folder3)

        individual_folder4 = os.path.join(individual_folder2, filename4)
        if not os.path.exists(individual_folder4):
            os.mkdir(individual_folder4)

        wb.save(os.path.join(individual_folder2, f"Checklist_{filename2}.xlsx"))

    # 생성된 모든 파일들을 하나의 압축 파일로 만듭니다.
    shutil.make_archive(output_folder, 'zip', output_folder)
    return f"{output_folder}.zip"

# afc checklist 자동생성 함수
def afc_checklist(uploaded_file, filename1):
    # 파일 로드
    grouped_data = pd.read_excel(uploaded_file, sheet_name=None)

    output_folder = "Checklist_Temp"
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)
    
    # 각 그룹(시트)에 대해 반복
    for sheet_name, sheet_data in grouped_data.items():
        # 새로운 체크리스트 워크북을 만듭니다 (위에서 제공한 코드를 계속 사용합니다).
        # ... [워크북 생성 및 서식 지정 코드] ...

        # Create a new Excel workbook and get the active worksheet
        wb = openpyxl.Workbook()
        ws = wb.active

        # Define the number of rows and columns
        num_rows = 45
        num_columns = 8  # Columns A to H

        # Populate the table with placeholder data
        for row in range(1, num_rows + 1):
            for col in range(1, num_columns + 1):
                cell = ws[get_column_letter(col) + str(row)]
                cell.value = None

        # 열 너비 설정
        column_widths = {'A': 14,'B': 35.5,'C': 7.9,'D': 29,'E': 12,'F': 28,'G': 21,'H': 29}

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # 행 높이 설정
        row_heights = {1: 15, 6: 15, 2: 37.8}

        # Apply the specified heights
        for row_num, height in row_heights.items():
            ws.row_dimensions[row_num].height = height

        # Set the height for all other rows
        for row_num in range(1, num_rows + 1):
            if row_num not in row_heights:
                ws.row_dimensions[row_num].height = 30

        # 2행 병합
        ws.merge_cells('A2:H2')

        # 2행 문자 삽입
        cell_2 = ws['A2']
        cell_2.value = "AFC 개방검사 검사표"
        cell_2.font = Font(bold=True, underline='single', size=25)
        cell_2.alignment = Alignment(horizontal='center', vertical='center')

        # 3~5행 병합
        ws.merge_cells('C3:E3')
        ws.merge_cells('C4:E4')
        ws.merge_cells('C5:E5')

        # 3행 문자 삽입
        cell_3_a = ws['A3']
        cell_3_a.value = '장치번호'
        cell_3_a.font = Font(bold=True, size=14)
        cell_3_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_3_f = ws['F3']
        cell_3_f.value = '장치등급'
        cell_3_f.font = Font(bold=True, size=14)
        cell_3_f.alignment = Alignment(horizontal='left', vertical='center')

        # 4행 문자 삽입
        cell_4_a = ws['A4']
        cell_4_a.value = '검사일 :'
        cell_4_a.font = Font(bold=True, size=14)
        cell_4_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_4_f = ws['F4']
        cell_4_f.value = '검사구분'
        cell_4_f.font = Font(bold=True, size=14)
        cell_4_f.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 문자 삽입
        cell_5_a = ws['A5']
        cell_5_a.value = '검사원 :'
        cell_5_a.font = Font(bold=True, size=14)
        cell_5_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_c = ws['C5']
        cell_5_c.value = '(서명)'
        cell_5_c.font = Font(bold=True, size=11)
        cell_5_c.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_f = ws['F5']
        cell_5_f.value = '엔지니어 :'
        cell_5_f.font = Font(bold=True, size=14)
        cell_5_f.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_h = ws['H5']
        cell_5_h.value = '(서명)'
        cell_5_h.font = Font(bold=True, size=11)
        cell_5_h.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 아래 굵은 테두리 표시
        bold_bottom_border = Border(bottom=Side(style='thick'))

        for col in range(1, 9):
            ws.cell(row=5, column=col).border = bold_bottom_border

        # 7행 병합 및 문자 삽입
        ws.merge_cells('A7:H7')
        cell_7_a = ws['A7']
        cell_7_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_7_a.value = '1. TA Worklist'
        cell_7_a.font = Font(bold=True, size=12)
        cell_7_a.alignment = Alignment(horizontal='left', vertical='center')

        # 8행 병합 및 문자 삽입
        ws.merge_cells('B8:E8')
        cell_8_a = ws['A8']
        cell_8_a.value = '번호'
        cell_8_a.font = Font(bold=True, size=12)
        cell_8_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('F8:H8')
        cell_8_f = ws['F8']
        cell_8_f.value = '주요 검사 사항'
        cell_8_f.font = Font(bold=True, size=12)
        cell_8_f.alignment = Alignment(horizontal='center', vertical='center')

        cell_8_b = ws['B8']
        cell_8_b.value = '작업 내용'
        cell_8_b.font = Font(bold=True, size=12)
        cell_8_b.alignment = Alignment(horizontal='center', vertical='center')

        # 9~18행 병합 및 문자 삽입
        for i in range(9, 19):
            ws.merge_cells(f'B{i}:E{i}')
            ws.merge_cells(f'F{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 8)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')      

        # 19행 병합 및 문자 삽입
        ws.merge_cells('A19:H19')
        cell_19_a = ws['A19']
        cell_19_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_19_a.value = '2. Damage Mechanism :'
        cell_19_a.font = Font(bold=True, size=12)
        cell_19_a.alignment = Alignment(horizontal='left', vertical='center')

        # 20행 병합 및 문자 삽입
        ws.merge_cells('A20:H20')
        cell_20_a = ws['A20']
        cell_20_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_20_a.value = '3. 장치 Type별 일반 검사 항목 (양호, 불량, N/A)'
        cell_20_a.font = Font(bold=True, size=12)
        cell_20_a.alignment = Alignment(horizontal='left', vertical='center')

        # 21행 병합 및 문자 삽입
        cell_21_a = ws['A21']
        cell_21_a.value = '구분'
        cell_21_a.font = Font(bold=True, size=12)
        cell_21_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_b = ws['B21']
        cell_21_b.value = '검사항목'
        cell_21_b.font = Font(bold=True, size=12)
        cell_21_b.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('C21:D21')
        cell_21_c = ws['C21']
        cell_21_c.value = '검사 결과'
        cell_21_c.font = Font(bold=True, size=12)
        cell_21_c.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_e = ws['E21']
        cell_21_e.value = '구분'
        cell_21_e.font = Font(bold=True, size=12)
        cell_21_e.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_f = ws['F21']
        cell_21_f.value = '검사항목'
        cell_21_f.font = Font(bold=True, size=12)
        cell_21_f.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('G21:H21')
        cell_21_g = ws['G21']
        cell_21_g.value = '검사 결과'
        cell_21_g.font = Font(bold=True, size=12)
        cell_21_g.alignment = Alignment(horizontal='center', vertical='center')

        # 22행 문자 삽입
        cell_22_b = ws['B22']
        cell_22_b.value = '청소 전 초기상태'
        cell_22_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_22_f= ws['F22']
        cell_22_f.value = 'Shell & Head(Header Box)'
        cell_22_f.alignment = Alignment(horizontal='center', vertical='center')

        # 23행 병합 및 문자 삽입
        cell_23_b = ws['B23']
        cell_23_b.value = 'Sludge, Scale 부착정도'
        cell_23_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_23_f= ws['F23']
        cell_23_f.value = 'Channel & Cover'
        cell_23_f.alignment = Alignment(horizontal='center', vertical='center')

        # 24행 병합 및 문자 삽입
        cell_24_b = ws['B24']
        cell_24_b.value = 'Shell & Head'
        cell_24_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_24_f= ws['F24']
        cell_24_f.value = 'Nozzle & Flange'
        cell_24_f.alignment = Alignment(horizontal='center', vertical='center')

        # 25행 문자 삽입
        cell_25_b = ws['B25']
        cell_25_b.value = 'Channel & Cover'
        cell_25_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_25_f= ws['F25']
        cell_25_f.value = 'IRIS/ECT (본)'
        cell_25_f.alignment = Alignment(horizontal='center', vertical='center')

        # 26행 문자 삽입
        cell_26_b = ws['B26']
        cell_26_b.value = 'Nozzle & Flange'
        cell_26_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_26_f= ws['F26']
        cell_26_f.value = 'Shell/Channel Main 용접심'
        cell_26_f.alignment = Alignment(horizontal='center', vertical='center')

        # 27행 문자 삽입
        cell_27_b = ws['B27']
        cell_27_b.value = 'Floating Head'
        cell_27_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_27_f = ws['F27']
        cell_27_f.value = 'Shell/Channel 부착물 용접심'
        cell_27_f.alignment = Alignment(horizontal='center', vertical='center')

        # 28행 문자 삽입
        cell_28_b = ws['B28']
        cell_28_b.value = 'Header Box'
        cell_28_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_28_f = ws['F28']
        cell_28_f.value = 'Shell/Channel to Nozzle 용접심'
        cell_28_f.alignment = Alignment(horizontal='center', vertical='center')

        # 29행 문자 삽입
        cell_29_b = ws['B29']
        cell_29_b.value = 'Header Box 나사산'
        cell_29_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_29_f= ws['F29']
        cell_29_f.value = 'Nozzle C & B Joint'
        cell_29_f.alignment = Alignment(horizontal='center', vertical='center')

        # 30행 문자 삽입
        cell_30_b = ws['B30']
        cell_30_b.value = 'Plug Bolt 나사산'
        cell_30_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_30_f= ws['F30']
        cell_30_f.value = 'U-Bend RT'
        cell_30_f.alignment = Alignment(horizontal='center', vertical='center')

        # 31행 문자 삽입
        cell_31_b = ws['B31']
        cell_31_b.value = 'Plug Gasket Face'
        cell_31_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_31_f= ws['F31']
        cell_31_f.value = '상부 1열 RT(Fixed Type)'
        cell_31_f.alignment = Alignment(horizontal='center', vertical='center')

        # 32행 문자 삽입
        cell_32_b = ws['B32']
        cell_32_b.value = 'Coating/Lining 상태'
        cell_32_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_32_f= ws['F32']
        cell_32_f.value = 'Strength Welding부 PT'
        cell_32_f.alignment = Alignment(horizontal='center', vertical='center')

        # 33행 문자 삽입
        cell_33_b = ws['B33']
        cell_33_b.value = 'Tube 외부'
        cell_33_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_33_f= ws['F33']
        cell_33_f.value = 'Header Box 용접심'
        cell_33_f.alignment = Alignment(horizontal='center', vertical='center')

        # 34행 문자 삽입
        cell_34_b = ws['B34']
        cell_34_b.value = 'Tube 내부'
        cell_34_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_34_f= ws['F34']
        cell_34_f.value = 'Partition/Stiffener Plate 용접부'
        cell_34_f.alignment = Alignment(horizontal='center', vertical='center')

        # 35행 문자 삽입
        cell_35_b = ws['B35']
        cell_35_b.value = 'Tube Sheet'
        cell_35_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_35_f= ws['F35']
        cell_35_f.value = '용접'
        cell_35_f.alignment = Alignment(horizontal='center', vertical='center')

        # 36행 문자 삽입
        cell_36_b = ws['B36']
        cell_36_b.value = 'Baffle, Tie-Rods'
        cell_36_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_36_f= ws['F36']
        cell_36_f.value = '열처리'
        cell_36_f.alignment = Alignment(horizontal='center', vertical='center')

        # 37행 문자 삽입
        cell_37_b = ws['B37']
        cell_37_b.value = 'Sealing Strip'
        cell_37_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_37_f = ws['F37']
        cell_37_f.value = 'Grinding'
        cell_37_f.alignment = Alignment(horizontal='center', vertical='center')

        # 38행 문자 삽입
        cell_38_b = ws['B38']
        cell_38_b.value = 'Impingement Plate'
        cell_38_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_38_f = ws['F38']
        cell_38_f.value = '보수부위 NDT'
        cell_38_f.alignment = Alignment(horizontal='center', vertical='center')

        # 39행 문자 삽입
        cell_39_b = ws['B39']
        cell_39_b.value = '보온재'
        cell_39_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_39_f = ws['F39']
        cell_39_f.value = '압력시험'
        cell_39_f.alignment = Alignment(horizontal='center', vertical='center')

        # 40행 문자 삽입
        cell_40_b = ws['B40']
        cell_40_b.value = 'Nozzle & Flange'
        cell_40_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_40_f = ws['F40']
        cell_40_f.value = 'Painting'
        cell_40_f.alignment = Alignment(horizontal='center', vertical='center')

        # 41행 문자 삽입
        cell_41_b = ws['B41']
        cell_41_b.value = 'Fittings'
        cell_41_b.alignment = Alignment(horizontal='center', vertical='center')

        # 42행 문자 삽입
        cell_42_b = ws['B42']
        cell_42_b.value = 'Supports'
        cell_42_b.alignment = Alignment(horizontal='center', vertical='center')

        # 43행 문자 삽입
        cell_43_b = ws['B43']
        cell_43_b.value = 'Fire Proofing'
        cell_43_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_43_f = ws['F43']
        cell_43_f.value = '1차'
        cell_43_f.alignment = Alignment(horizontal='center', vertical='center')

        # 44행 문자 삽입
        cell_44_b = ws['B44']
        cell_44_b.value = 'Paint 상태'
        cell_44_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_44_f = ws['F44']
        cell_44_f.value = '2차'
        cell_44_f.alignment = Alignment(horizontal='center', vertical='center')

        # 45행 문자 삽입
        cell_45_b = ws['B45']
        cell_45_b.value = 'Floating Head Anchor bolt\nDe-bolting 유무 점검'
        cell_45_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_f = ws['F45']
        cell_45_f.value = '3차'
        cell_45_f.alignment = Alignment(horizontal='center', vertical='center')

        # 세로 병합
        ws.merge_cells('A22:A32')
        cell_22_a = ws['A22']
        cell_22_a.value = '내부\n육안검사'
        cell_22_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E22:E25')
        cell_22_e = ws['E22']
        cell_22_e.value = '두께측정'
        cell_22_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A33:A38')
        cell_33_a = ws['A33']
        cell_33_a.value = 'Bundle\n육안검사'
        cell_33_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E26:E34')
        cell_26_e = ws['E26']
        cell_26_e.value = '내부\nNDT'
        cell_26_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A39:A45')
        cell_39_a = ws['A39']
        cell_39_a.value = '외부 육안\n검사'
        cell_39_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E35:E42')
        cell_35_e = ws['E35']
        cell_35_e.value = '추가\n보수작업'
        cell_35_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E43:E45')
        cell_43_e = ws['E43']
        cell_43_e.value = '수압시험'
        cell_43_e.alignment = Alignment(horizontal='center', vertical='center')

        # 테두리 적용
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))

        # Apply the border to all cells in the range
        for row in ws.iter_rows(min_row=1, max_row=num_rows, min_col=1, max_col=num_columns):
            for cell in row:
                cell.border = thin_border
      ########################################################################### 여기부터 수정 필요           
    # =============================================================================================================

        # 장치 번호 입력
        device_number = sheet_data.iloc[0, 0]  # 장치 번호 추정 위치
        ws['B3'].value = device_number  # 장치 번호 셀 할당
        ws['B3'].alignment = Alignment(horizontal='center', vertical='center')

    # =============================================================================================================    

        # 작업 사항 및 검사 사항 사이의 데이터를 처리하기 위한 로직
        worklist_start = None
        content_end = None  # 작업 사항과 검사 사항 사이의 내용을 저장할 위치
        separate_row_found = False  # '띄어쓰기' 행이 발견되었는지 확인

        for i, row in sheet_data.iterrows():
            # '[작업사항]'이 발견되면 시작 위치를 설정합니다.
            if '[작업사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                worklist_start = i

            # '[검사사항]'이 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '[검사사항]' in str(row[0]) and worklist_start is not None:
                content_end = i
                break  # '[검사사항]'이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

            # 띄어쓰기가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif row[0] == '띄어쓰기' and worklist_start is not None:  # row[0]이 빈 문자열인 경우를 확인합니다.
                content_end = i
                empty_row_found = True  # 빈 행이 발견되었음을 표시합니다.
                break  # 빈 행이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 작업 사항과 검사 사항 사이에 데이터가 있고, '[검사사항]' 또는 빈 행이 발견된 경우 데이터 삽입을 시작합니다.
        if worklist_start is not None and content_end is not None:
            worklist_data = sheet_data.iloc[worklist_start + 1:content_end]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(worklist_data[first_column_name], start=9):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

    # =============================================================================================================  
        # 검사 사항과 '띄어쓰기' 사이의 데이터를 처리하기 위한 로직
        inpection_start = None
        content_end1 = None  # 검사 사항과 띄어쓰기 사이의 내용을 저장할 위치

        for i, row in sheet_data.iterrows():
            # '[검사사항]'이 발견되면 시작 위치를 설정합니다.
            if '[검사사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                inpection_start = i

            # '띄어쓰기'가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '띄어쓰기' in str(row[0]) and inpection_start is not None:
                content_end1 = i
                break  # '띄어쓰기'가 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 검사 사항과 '띄어쓰기' 사이에 데이터 삽입을 시작합니다.
        if inpection_start is not None and content_end1 is not None:
            inpection_data = sheet_data.iloc[inpection_start + 1:content_end1]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name1 = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # F9셀부터 데이터 삽입
            for i, item in enumerate(inpection_data[first_column_name1], start=9):
                cell = ws.cell(row=i, column=6)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기


        # 파일을 지정된 폴더에 저장합니다.

        # 각 파일에 대한 개별 폴더 생성
        filename1 = filename1
        filename2 = device_number.replace("/", "&").strip()
        filename3 = 'Photo'
        filename4 = '도면'

        individual_folder1 = os.path.join(output_folder, filename1)
        if not os.path.exists(individual_folder1):
            os.mkdir(individual_folder1)

        individual_folder2 = os.path.join(individual_folder1, filename2)
        if not os.path.exists(individual_folder2):
            os.mkdir(individual_folder2)

        individual_folder3 = os.path.join(individual_folder2, filename3)
        if not os.path.exists(individual_folder3):
            os.mkdir(individual_folder3)

        individual_folder4 = os.path.join(individual_folder2, filename4)
        if not os.path.exists(individual_folder4):
            os.mkdir(individual_folder4)

        wb.save(os.path.join(individual_folder2, f"Checklist_{filename2}.xlsx"))

    # 생성된 모든 파일들을 하나의 압축 파일로 만듭니다.
    shutil.make_archive(output_folder, 'zip', output_folder)
    return f"{output_folder}.zip"

# cdv checklist 통합본 생성 함수
def cdv_checklist_merge(uploaded_file):
    # 파일 로드
    grouped_data = pd.read_excel(uploaded_file, sheet_name=None)

    output_folder = "Checklist_merged_Temp"
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)

    # 새로운 Excel 워크북을 생성합니다.
    wb = openpyxl.Workbook()

    # 기본 제공되는 빈 워크시트 제거
    wb.remove(wb.active)
    
    for sheet_name, sheet_data in grouped_data.items():
        # 새 워크시트를 생성하고 이름을 설정합니다.
        ws = wb.create_sheet(title=sheet_name)

        # Define the number of rows and columns
        num_rows = 66
        num_columns = 8  # Columns A to H

        # Populate the table with placeholder data
        for row in range(1, num_rows + 1):
            for col in range(1, num_columns + 1):
                cell = ws[get_column_letter(col) + str(row)]
                cell.value = None

        # 열 너비 설정
        column_widths = {'A': 9.9,'B': 24.3,'C': 7.9,'D': 29.9,'E': 9.9,'F': 17.5,'G': 7.9,'H': 29.9}

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # 행 높이 설정
        row_heights = {1: 15, 6: 15, 2: 37.8}

        # Apply the specified heights
        for row_num, height in row_heights.items():
            ws.row_dimensions[row_num].height = height

        # Set the height for all other rows
        for row_num in range(1, num_rows + 1):
            if row_num not in row_heights:
                ws.row_dimensions[row_num].height = 30

        # 2행 병합
        ws.merge_cells('A2:H2')

        # 2행 문자 삽입
        cell_2 = ws['A2']
        cell_2.value = "압력용기 개방검사 검사표"
        cell_2.font = Font(bold=True, underline='single', size=25)
        cell_2.alignment = Alignment(horizontal='center', vertical='center')

        # 3~5행 병합
        ws.merge_cells('C3:E3')
        ws.merge_cells('C4:E4')
        ws.merge_cells('C5:E5')

        # 3행 문자 삽입
        cell_3_a = ws['A3']
        cell_3_a.value = '장치번호'
        cell_3_a.font = Font(bold=True, size=14)
        cell_3_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_3_f = ws['F3']
        cell_3_f.value = '장치등급'
        cell_3_f.font = Font(bold=True, size=14)
        cell_3_f.alignment = Alignment(horizontal='left', vertical='center')

        # 4행 문자 삽입
        cell_4_a = ws['A4']
        cell_4_a.value = '검사일 :'
        cell_4_a.font = Font(bold=True, size=14)
        cell_4_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_4_f = ws['F4']
        cell_4_f.value = '검사구분'
        cell_4_f.font = Font(bold=True, size=14)
        cell_4_f.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 문자 삽입
        cell_5_a = ws['A5']
        cell_5_a.value = '검사원 :'
        cell_5_a.font = Font(bold=True, size=14)
        cell_5_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_c = ws['C5']
        cell_5_c.value = '(서명)'
        cell_5_c.font = Font(bold=True, size=11)
        cell_5_c.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_f = ws['F5']
        cell_5_f.value = '엔지니어 :'
        cell_5_f.font = Font(bold=True, size=14)
        cell_5_f.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_h = ws['H5']
        cell_5_h.value = '(서명)'
        cell_5_h.font = Font(bold=True, size=11)
        cell_5_h.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 아래 굵은 테두리 표시
        bold_bottom_border = Border(bottom=Side(style='thick'))

        for col in range(1, 9):
            ws.cell(row=5, column=col).border = bold_bottom_border

        # 7행 병합 및 문자 삽입
        ws.merge_cells('A7:H7')
        cell_7_a = ws['A7']
        cell_7_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_7_a.value = '1. TA Worklist'
        cell_7_a.font = Font(bold=True, size=12)
        cell_7_a.alignment = Alignment(horizontal='left', vertical='center')

        # 8행 병합 및 문자 삽입
        ws.merge_cells('B8:H8')
        cell_8_a = ws['A8']
        cell_8_a.value = '번호'
        cell_8_a.font = Font(bold=True, size=12)
        cell_8_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_8_b = ws['B8']
        cell_8_b.value = '작업 사항'
        cell_8_b.font = Font(bold=True, size=12)
        cell_8_b.alignment = Alignment(horizontal='center', vertical='center')

        # 9~18행 병합 및 문자 삽입
        for i in range(9, 19):
            ws.merge_cells(f'B{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 8)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 19행 병합 및 문자 삽입
        ws.merge_cells('A19:H19')
        cell_19_a = ws['A19']
        cell_19_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_19_a.value = '2. Damage Mechanism :'
        cell_19_a.font = Font(bold=True, size=12)
        cell_19_a.alignment = Alignment(horizontal='left', vertical='center')

        # 20행 병합 및 문자 삽입
        ws.merge_cells('A20:H20')
        cell_20_a = ws['A20']
        cell_20_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_20_a.value = '3. 중요 검사 항목'
        cell_20_a.font = Font(bold=True, size=12)
        cell_20_a.alignment = Alignment(horizontal='left', vertical='center')

        # 21행 병합 및 문자 삽입
        ws.merge_cells('B21:H21')
        cell_21_a = ws['A21']
        cell_21_a.value = '번호'
        cell_21_a.font = Font(bold=True, size=12)
        cell_21_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_b = ws['B21']
        cell_21_b.value = '검사 항목'
        cell_21_b.font = Font(bold=True, size=12)
        cell_21_b.alignment = Alignment(horizontal='center', vertical='center')

        # 22행~31행 병합 및 문자 삽입
        for i in range(22, 32):
            ws.merge_cells(f'B{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 21)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 32행 병합 및 문자 삽입
        ws.merge_cells('A32:H32')
        cell_32_a = ws['A32']
        cell_32_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_32_a.value = '4. 주요 검사 이력 (전기 TA, 운전 중 이력)'
        cell_32_a.font = Font(bold=True, size=12)
        cell_32_a.alignment = Alignment(horizontal='left', vertical='center')

        # 33행 병합 및 문자 삽입
        ws.merge_cells('B33:H33')
        cell_33_a = ws['A33']
        cell_33_a.value = '번호'
        cell_33_a.font = Font(bold=True, size=12)
        cell_33_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_33_b = ws['B33']
        cell_33_b.value = '검사 이력'
        cell_33_b.font = Font(bold=True, size=12)
        cell_33_b.alignment = Alignment(horizontal='center', vertical='center')

        # 34행~43행 병합 및 문자 삽입
        for i in range(34, 44):
            ws.merge_cells(f'B{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 33)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 44행 병합 및 문자 삽입
        ws.merge_cells('A44:H44')
        cell_44_a = ws['A44']
        cell_44_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_44_a.value = '5. 장치 Type별 일반 검사 항목 (양호, 불량, N/A)'
        cell_44_a.font = Font(bold=True, size=12)
        cell_44_a.alignment = Alignment(horizontal='left', vertical='center')

        # 45행 병합 및 문자 삽입
        cell_45_a = ws['A45']
        cell_45_a.value = '구분'
        cell_45_a.font = Font(bold=True, size=12)
        cell_45_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_b = ws['B45']
        cell_45_b.value = '검사항목'
        cell_45_b.font = Font(bold=True, size=12)
        cell_45_b.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('C45:D45')
        cell_45_c = ws['C45']
        cell_45_c.value = '검사 결과'
        cell_45_c.font = Font(bold=True, size=12)
        cell_45_c.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_e = ws['E45']
        cell_45_e.value = '구분'
        cell_45_e.font = Font(bold=True, size=12)
        cell_45_e.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_f = ws['F45']
        cell_45_f.value = '검사항목'
        cell_45_f.font = Font(bold=True, size=12)
        cell_45_f.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('G45:H45')
        cell_45_g = ws['G45']
        cell_45_g.value = '검사 결과'
        cell_45_g.font = Font(bold=True, size=12)
        cell_45_g.alignment = Alignment(horizontal='center', vertical='center')

        # 46행 문자 삽입
        cell_46_b = ws['B46']
        cell_46_b.value = '청소 전 초기상태'
        cell_46_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_46_f= ws['F46']
        cell_46_f.value = 'Tray Support'
        cell_46_f.alignment = Alignment(horizontal='center', vertical='center')

        # 47행 병합 및 문자 삽입
        cell_47_b = ws['B47']
        cell_47_b.value = 'Sludge, Scale 부착정도'
        cell_47_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_47_f= ws['F47']
        cell_47_f.value = 'Down Comer'
        cell_47_f.alignment = Alignment(horizontal='center', vertical='center')

        # 문자 삽입
        cell_48_b = ws['B48']
        cell_48_b.value = 'Shell'
        cell_48_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_48_f= ws['F48']
        cell_48_f.value = 'Draw Off'
        cell_48_f.alignment = Alignment(horizontal='center', vertical='center')

        # 49 행 문자 삽입
        cell_49_b = ws['B49']
        cell_49_b.value = 'Head'
        cell_49_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_49_f= ws['F49']
        cell_49_f.value = 'Entry Horn'
        cell_49_f.alignment = Alignment(horizontal='center', vertical='center')

        # 50행 문자 삽입
        cell_50_b = ws['B50']
        cell_50_b.value = 'Boot'
        cell_50_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_50_f= ws['F50']
        cell_50_f.value = 'Packing'
        cell_50_f.alignment = Alignment(horizontal='center', vertical='center')

        # 51행 문자 삽입
        cell_51_b = ws['B51']
        cell_51_b.value = 'Nozzle & Flange'
        cell_51_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_51_f= ws['F51']
        cell_51_f.value = 'Electrode'
        cell_51_f.alignment = Alignment(horizontal='center', vertical='center')

        # 52행 문자 삽입
        cell_52_b = ws['B52']
        cell_52_b.value = 'Lining/Paint 상태'
        cell_52_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_52_f= ws['F52']
        cell_52_f.value = 'Grinding'
        cell_52_f.alignment = Alignment(horizontal='center', vertical='center')

        # 53행 문자 삽입
        cell_53_b = ws['B53']
        cell_53_b.value = 'Main 용접심'
        cell_53_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_53_f= ws['F53']
        cell_53_f.value = '용접'
        cell_53_f.alignment = Alignment(horizontal='center', vertical='center')

        # 54행 문자 삽입
        cell_54_b = ws['B54']
        cell_54_b.value = '부착물 용접심'
        cell_54_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_54_f= ws['F54']
        cell_54_f.value = '열처리'
        cell_54_f.alignment = Alignment(horizontal='center', vertical='center')

        # 55행 문자 삽입
        cell_55_b = ws['B55']
        cell_55_b.value = 'Shell to Nozzle 용접심'
        cell_55_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_55_f= ws['F55']
        cell_55_f.value = '보수부위 NDT'
        cell_55_f.alignment = Alignment(horizontal='center', vertical='center')

        # 56행 문자 삽입
        cell_56_b = ws['B56']
        cell_56_b.value = '노즐 Sleeve'
        cell_56_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_56_f= ws['F56']
        cell_56_f.value = '압력시험'
        cell_56_f.alignment = Alignment(horizontal='center', vertical='center')

        # 57행 문자 삽입
        cell_57_b = ws['B57']
        cell_57_b.value = 'Distributor / Collector'
        cell_57_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_57_f= ws['F57']
        cell_57_f.value = 'Painting'
        cell_57_f.alignment = Alignment(horizontal='center', vertical='center')

        # 58행 문자 삽입
        cell_58_b = ws['B58']
        cell_58_b.value = 'Internal Piping'
        cell_58_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_58_f= ws['F58']
        cell_58_f.value = 'Tray M/W'
        cell_58_f.alignment = Alignment(horizontal='center', vertical='center')

        # 59행 문자 삽입
        cell_59_b = ws['B59']
        cell_59_b.value = 'Johnson Screen'
        cell_59_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_59_f= ws['F59']
        cell_59_f.value = 'Distributor Pipe'
        cell_59_f.alignment = Alignment(horizontal='center', vertical='center')

        # 60행 문자 삽입
        cell_60_b = ws['B60']
        cell_60_b.value = 'Wire Mesh'
        cell_60_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_60_f= ws['F60']
        cell_60_f.value = 'Shell/Head'
        cell_60_f.alignment = Alignment(horizontal='center', vertical='center')

        # 61행 문자 삽입
        cell_61_b = ws['B61']
        cell_61_b.value = 'Grid'
        cell_61_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_61_f= ws['F61']
        cell_61_f.value = 'Nozzle'
        cell_61_f.alignment = Alignment(horizontal='center', vertical='center')

        # 62행 문자 삽입
        cell_62_b = ws['B62']
        cell_62_b.value = 'Screen / Grid Support'
        cell_62_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_62_f= ws['F62']
        cell_62_f.value = '보온재'
        cell_62_f.alignment = Alignment(horizontal='center', vertical='center')

        # 63행 문자 삽입
        cell_63_b = ws['B63']
        cell_63_b.value = 'Vortex Breaker'
        cell_63_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_63_f= ws['F63']
        cell_63_f.value = 'Nozzle & Flange'
        cell_63_f.alignment = Alignment(horizontal='center', vertical='center')

        # 64행 문자 삽입
        cell_64_b = ws['B64']
        cell_64_b.value = 'Tray'
        cell_64_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_64_f= ws['F64']
        cell_64_f.value = 'Fittings'
        cell_64_f.alignment = Alignment(horizontal='center', vertical='center')

        # 65행 문자 삽입
        cell_65_b = ws['B65']
        cell_65_b.value = 'Tray Cap'
        cell_65_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_65_f= ws['F65']
        cell_65_f.value = 'Supports'
        cell_65_f.alignment = Alignment(horizontal='center', vertical='center')

        # 66행 문자 삽입
        cell_66_b = ws['B66']
        cell_66_b.value = 'Tray Bolt/Nut/Clamp'
        cell_66_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_66_f= ws['F66']
        cell_66_f.value = 'Fire Proofing'
        cell_66_f.alignment = Alignment(horizontal='center', vertical='center')

        # 세로 병합
        ws.merge_cells('A46:A52')
        cell_46_a = ws['A46']
        cell_46_a.value = '내부\n육안검사'
        cell_46_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E46:E51')
        cell_46_e = ws['E46']
        cell_46_e.value = 'Internals\n상태검사'
        cell_46_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A53:A57')
        cell_53_a = ws['A53']
        cell_53_a.value = '내부 NDT'
        cell_53_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E52:E57')
        cell_52_e = ws['E52']
        cell_52_e.value = '추가\n보수작업'
        cell_52_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A58:A66')
        cell_58_a = ws['A58']
        cell_58_a.value = 'Internals\n상태검사'
        cell_58_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E58:E59')
        cell_58_e = ws['E58']
        cell_58_e.value = 'Internal\n조립'
        cell_58_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E60:E61')
        cell_60_e = ws['E60']
        cell_60_e.value = '두께측정'
        cell_60_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E62:E66')
        cell_62_e = ws['E62']
        cell_62_e.value = '외부\n육안 검사'
        cell_62_e.alignment = Alignment(horizontal='center', vertical='center')

        # 테두리 적용
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))

        # Apply the border to all cells in the range
        for row in ws.iter_rows(min_row=1, max_row=num_rows, min_col=1, max_col=num_columns):
            for cell in row:
                cell.border = thin_border

    # =============================================================================================================
        # 장치 번호 입력
        device_number = sheet_data.iloc[0, 0]  # 장치 번호 추정 위치
        ws['B3'].value = device_number  # 장치 번호 셀 할당
        ws['B3'].alignment = Alignment(horizontal='center', vertical='center')

    # =============================================================================================================    

        # 작업 사항 및 검사 사항 사이의 데이터를 처리하기 위한 로직
        worklist_start = None
        content_end = None  # 작업 사항과 검사 사항 사이의 내용을 저장할 위치
        separate_row_found = False  # '띄어쓰기' 행이 발견되었는지 확인

        for i, row in sheet_data.iterrows():
            # '[작업사항]'이 발견되면 시작 위치를 설정합니다.
            if '[작업사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                worklist_start = i

            # '[검사사항]'이 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '[검사사항]' in str(row[0]) and worklist_start is not None:
                content_end = i
                break  # '[검사사항]'이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

            # 띄어쓰기가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif row[0] == '띄어쓰기' and worklist_start is not None:  # row[0]이 빈 문자열인 경우를 확인합니다.
                content_end = i
                empty_row_found = True  # 빈 행이 발견되었음을 표시합니다.
                break  # 빈 행이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 작업 사항과 검사 사항 사이에 데이터가 있고, '[검사사항]' 또는 빈 행이 발견된 경우 데이터 삽입을 시작합니다.
        if worklist_start is not None and content_end is not None:
            worklist_data = sheet_data.iloc[worklist_start + 1:content_end]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(worklist_data[first_column_name], start=9):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

    # =============================================================================================================  
        # 검사 사항과 '띄어쓰기' 사이의 데이터를 처리하기 위한 로직
        inpection_start = None
        content_end1 = None  # 검사 사항과 띄어쓰기 사이의 내용을 저장할 위치

        for i, row in sheet_data.iterrows():
            # '[검사사항]'이 발견되면 시작 위치를 설정합니다.
            if '[검사사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                inpection_start = i

            # '띄어쓰기'가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '띄어쓰기' in str(row[0]) and inpection_start is not None:
                content_end1 = i
                break  # '띄어쓰기'가 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 검사 사항과 '띄어쓰기' 사이에 데이터 삽입을 시작합니다.
        if inpection_start is not None and content_end1 is not None:
            inpection_data = sheet_data.iloc[inpection_start + 1:content_end1]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name1 = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(inpection_data[first_column_name1], start=22):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

        # 워크북을 저장합니다. 파일 이름은 원하는 대로 설정할 수 있습니다.
        output_excel_file = "CDV,HTR_Checklist 통합.xlsx"
        saved_file_path = os.path.join(output_folder, output_excel_file)
        wb.save(saved_file_path)
    
    return saved_file_path

# hex checklist 통합본 생성 함수
def hex_checklist_merge(uploaded_file):
    # 파일 로드
    grouped_data = pd.read_excel(uploaded_file, sheet_name=None)

    output_folder = "Checklist_merged_Temp"
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)

    # 새로운 Excel 워크북을 생성합니다.
    wb = openpyxl.Workbook()

    # 기본 제공되는 빈 워크시트 제거
    wb.remove(wb.active)
    
    for sheet_name, sheet_data in grouped_data.items():
        # 새 워크시트를 생성하고 이름을 설정합니다.
        ws = wb.create_sheet(title=sheet_name)

        # Define the number of rows and columns
        num_rows = 45
        num_columns = 8  # Columns A to H

        # Populate the table with placeholder data
        for row in range(1, num_rows + 1):
            for col in range(1, num_columns + 1):
                cell = ws[get_column_letter(col) + str(row)]
                cell.value = None

        # 열 너비 설정
        column_widths = {'A': 14,'B': 35.5,'C': 7.9,'D': 29,'E': 12,'F': 28,'G': 21,'H': 29}

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # 행 높이 설정
        row_heights = {1: 15, 6: 15, 2: 37.8}

        # Apply the specified heights
        for row_num, height in row_heights.items():
            ws.row_dimensions[row_num].height = height

        # Set the height for all other rows
        for row_num in range(1, num_rows + 1):
            if row_num not in row_heights:
                ws.row_dimensions[row_num].height = 30

        # 2행 병합
        ws.merge_cells('A2:H2')

        # 2행 문자 삽입
        cell_2 = ws['A2']
        cell_2.value = "열교환기 개방검사 검사표"
        cell_2.font = Font(bold=True, underline='single', size=25)
        cell_2.alignment = Alignment(horizontal='center', vertical='center')

        # 3~5행 병합
        ws.merge_cells('C3:E3')
        ws.merge_cells('C4:E4')
        ws.merge_cells('C5:E5')

        # 3행 문자 삽입
        cell_3_a = ws['A3']
        cell_3_a.value = '장치번호'
        cell_3_a.font = Font(bold=True, size=14)
        cell_3_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_3_f = ws['F3']
        cell_3_f.value = '장치등급'
        cell_3_f.font = Font(bold=True, size=14)
        cell_3_f.alignment = Alignment(horizontal='left', vertical='center')

        # 4행 문자 삽입
        cell_4_a = ws['A4']
        cell_4_a.value = '검사일 :'
        cell_4_a.font = Font(bold=True, size=14)
        cell_4_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_4_f = ws['F4']
        cell_4_f.value = '검사구분'
        cell_4_f.font = Font(bold=True, size=14)
        cell_4_f.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 문자 삽입
        cell_5_a = ws['A5']
        cell_5_a.value = '검사원 :'
        cell_5_a.font = Font(bold=True, size=14)
        cell_5_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_c = ws['C5']
        cell_5_c.value = '(서명)'
        cell_5_c.font = Font(bold=True, size=11)
        cell_5_c.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_f = ws['F5']
        cell_5_f.value = '엔지니어 :'
        cell_5_f.font = Font(bold=True, size=14)
        cell_5_f.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_h = ws['H5']
        cell_5_h.value = '(서명)'
        cell_5_h.font = Font(bold=True, size=11)
        cell_5_h.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 아래 굵은 테두리 표시
        bold_bottom_border = Border(bottom=Side(style='thick'))

        for col in range(1, 9):
            ws.cell(row=5, column=col).border = bold_bottom_border

        # 7행 병합 및 문자 삽입
        ws.merge_cells('A7:H7')
        cell_7_a = ws['A7']
        cell_7_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_7_a.value = '1. TA Worklist'
        cell_7_a.font = Font(bold=True, size=12)
        cell_7_a.alignment = Alignment(horizontal='left', vertical='center')

        # 8행 병합 및 문자 삽입
        ws.merge_cells('B8:E8')
        cell_8_a = ws['A8']
        cell_8_a.value = '번호'
        cell_8_a.font = Font(bold=True, size=12)
        cell_8_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('F8:H8')
        cell_8_f = ws['F8']
        cell_8_f.value = '주요 검사 사항'
        cell_8_f.font = Font(bold=True, size=12)
        cell_8_f.alignment = Alignment(horizontal='center', vertical='center')

        cell_8_b = ws['B8']
        cell_8_b.value = '작업 내용'
        cell_8_b.font = Font(bold=True, size=12)
        cell_8_b.alignment = Alignment(horizontal='center', vertical='center')

        # 9~18행 병합 및 문자 삽입
        for i in range(9, 19):
            ws.merge_cells(f'B{i}:E{i}')
            ws.merge_cells(f'F{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 8)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')      

        # 19행 병합 및 문자 삽입
        ws.merge_cells('A19:H19')
        cell_19_a = ws['A19']
        cell_19_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_19_a.value = '2. Damage Mechanism :'
        cell_19_a.font = Font(bold=True, size=12)
        cell_19_a.alignment = Alignment(horizontal='left', vertical='center')

        # 20행 병합 및 문자 삽입
        ws.merge_cells('A20:H20')
        cell_20_a = ws['A20']
        cell_20_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_20_a.value = '3. 장치 Type별 일반 검사 항목 (양호, 불량, N/A)'
        cell_20_a.font = Font(bold=True, size=12)
        cell_20_a.alignment = Alignment(horizontal='left', vertical='center')

        # 21행 병합 및 문자 삽입
        cell_21_a = ws['A21']
        cell_21_a.value = '구분'
        cell_21_a.font = Font(bold=True, size=12)
        cell_21_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_b = ws['B21']
        cell_21_b.value = '검사항목'
        cell_21_b.font = Font(bold=True, size=12)
        cell_21_b.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('C21:D21')
        cell_21_c = ws['C21']
        cell_21_c.value = '검사 결과'
        cell_21_c.font = Font(bold=True, size=12)
        cell_21_c.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_e = ws['E21']
        cell_21_e.value = '구분'
        cell_21_e.font = Font(bold=True, size=12)
        cell_21_e.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_f = ws['F21']
        cell_21_f.value = '검사항목'
        cell_21_f.font = Font(bold=True, size=12)
        cell_21_f.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('G21:H21')
        cell_21_g = ws['G21']
        cell_21_g.value = '검사 결과'
        cell_21_g.font = Font(bold=True, size=12)
        cell_21_g.alignment = Alignment(horizontal='center', vertical='center')

        # 22행 문자 삽입
        cell_22_b = ws['B22']
        cell_22_b.value = '청소 전 초기상태'
        cell_22_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_22_f= ws['F22']
        cell_22_f.value = 'Shell & Head(Header Box)'
        cell_22_f.alignment = Alignment(horizontal='center', vertical='center')

        # 23행 병합 및 문자 삽입
        cell_23_b = ws['B23']
        cell_23_b.value = 'Sludge, Scale 부착정도'
        cell_23_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_23_f= ws['F23']
        cell_23_f.value = 'Channel & Cover'
        cell_23_f.alignment = Alignment(horizontal='center', vertical='center')

        # 24행 병합 및 문자 삽입
        cell_24_b = ws['B24']
        cell_24_b.value = 'Shell & Head'
        cell_24_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_24_f= ws['F24']
        cell_24_f.value = 'Nozzle & Flange'
        cell_24_f.alignment = Alignment(horizontal='center', vertical='center')

        # 25행 문자 삽입
        cell_25_b = ws['B25']
        cell_25_b.value = 'Channel & Cover'
        cell_25_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_25_f= ws['F25']
        cell_25_f.value = 'IRIS/ECT (본)'
        cell_25_f.alignment = Alignment(horizontal='center', vertical='center')

        # 26행 문자 삽입
        cell_26_b = ws['B26']
        cell_26_b.value = 'Nozzle & Flange'
        cell_26_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_26_f= ws['F26']
        cell_26_f.value = 'Shell/Channel Main 용접심'
        cell_26_f.alignment = Alignment(horizontal='center', vertical='center')

        # 27행 문자 삽입
        cell_27_b = ws['B27']
        cell_27_b.value = 'Floating Head'
        cell_27_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_27_f = ws['F27']
        cell_27_f.value = 'Shell/Channel 부착물 용접심'
        cell_27_f.alignment = Alignment(horizontal='center', vertical='center')

        # 28행 문자 삽입
        cell_28_b = ws['B28']
        cell_28_b.value = 'Header Box'
        cell_28_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_28_f = ws['F28']
        cell_28_f.value = 'Shell/Channel to Nozzle 용접심'
        cell_28_f.alignment = Alignment(horizontal='center', vertical='center')

        # 29행 문자 삽입
        cell_29_b = ws['B29']
        cell_29_b.value = 'Header Box 나사산'
        cell_29_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_29_f= ws['F29']
        cell_29_f.value = 'Nozzle C & B Joint'
        cell_29_f.alignment = Alignment(horizontal='center', vertical='center')

        # 30행 문자 삽입
        cell_30_b = ws['B30']
        cell_30_b.value = 'Plug Bolt 나사산'
        cell_30_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_30_f= ws['F30']
        cell_30_f.value = 'U-Bend RT'
        cell_30_f.alignment = Alignment(horizontal='center', vertical='center')

        # 31행 문자 삽입
        cell_31_b = ws['B31']
        cell_31_b.value = 'Plug Gasket Face'
        cell_31_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_31_f= ws['F31']
        cell_31_f.value = '상부 1열 RT(Fixed Type)'
        cell_31_f.alignment = Alignment(horizontal='center', vertical='center')

        # 32행 문자 삽입
        cell_32_b = ws['B32']
        cell_32_b.value = 'Coating/Lining 상태'
        cell_32_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_32_f= ws['F32']
        cell_32_f.value = 'Strength Welding부 PT'
        cell_32_f.alignment = Alignment(horizontal='center', vertical='center')

        # 33행 문자 삽입
        cell_33_b = ws['B33']
        cell_33_b.value = 'Tube 외부'
        cell_33_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_33_f= ws['F33']
        cell_33_f.value = 'Header Box 용접심'
        cell_33_f.alignment = Alignment(horizontal='center', vertical='center')

        # 34행 문자 삽입
        cell_34_b = ws['B34']
        cell_34_b.value = 'Tube 내부'
        cell_34_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_34_f= ws['F34']
        cell_34_f.value = 'Partition/Stiffener Plate 용접부'
        cell_34_f.alignment = Alignment(horizontal='center', vertical='center')

        # 35행 문자 삽입
        cell_35_b = ws['B35']
        cell_35_b.value = 'Tube Sheet'
        cell_35_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_35_f= ws['F35']
        cell_35_f.value = '용접'
        cell_35_f.alignment = Alignment(horizontal='center', vertical='center')

        # 36행 문자 삽입
        cell_36_b = ws['B36']
        cell_36_b.value = 'Baffle, Tie-Rods'
        cell_36_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_36_f= ws['F36']
        cell_36_f.value = '열처리'
        cell_36_f.alignment = Alignment(horizontal='center', vertical='center')

        # 37행 문자 삽입
        cell_37_b = ws['B37']
        cell_37_b.value = 'Sealing Strip'
        cell_37_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_37_f = ws['F37']
        cell_37_f.value = 'Grinding'
        cell_37_f.alignment = Alignment(horizontal='center', vertical='center')

        # 38행 문자 삽입
        cell_38_b = ws['B38']
        cell_38_b.value = 'Impingement Plate'
        cell_38_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_38_f = ws['F38']
        cell_38_f.value = '보수부위 NDT'
        cell_38_f.alignment = Alignment(horizontal='center', vertical='center')

        # 39행 문자 삽입
        cell_39_b = ws['B39']
        cell_39_b.value = '보온재'
        cell_39_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_39_f = ws['F39']
        cell_39_f.value = '압력시험'
        cell_39_f.alignment = Alignment(horizontal='center', vertical='center')

        # 40행 문자 삽입
        cell_40_b = ws['B40']
        cell_40_b.value = 'Nozzle & Flange'
        cell_40_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_40_f = ws['F40']
        cell_40_f.value = 'Painting'
        cell_40_f.alignment = Alignment(horizontal='center', vertical='center')

        # 41행 문자 삽입
        cell_41_b = ws['B41']
        cell_41_b.value = 'Fittings'
        cell_41_b.alignment = Alignment(horizontal='center', vertical='center')

        # 42행 문자 삽입
        cell_42_b = ws['B42']
        cell_42_b.value = 'Supports'
        cell_42_b.alignment = Alignment(horizontal='center', vertical='center')

        # 43행 문자 삽입
        cell_43_b = ws['B43']
        cell_43_b.value = 'Fire Proofing'
        cell_43_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_43_f = ws['F43']
        cell_43_f.value = '1차'
        cell_43_f.alignment = Alignment(horizontal='center', vertical='center')

        # 44행 문자 삽입
        cell_44_b = ws['B44']
        cell_44_b.value = 'Paint 상태'
        cell_44_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_44_f = ws['F44']
        cell_44_f.value = '2차'
        cell_44_f.alignment = Alignment(horizontal='center', vertical='center')

        # 45행 문자 삽입
        cell_45_b = ws['B45']
        cell_45_b.value = 'Saddle Sliding Side\nDe-Bolting 유무 점검'
        cell_45_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_f = ws['F45']
        cell_45_f.value = '3차'
        cell_45_f.alignment = Alignment(horizontal='center', vertical='center')

        # 세로 병합
        ws.merge_cells('A22:A32')
        cell_22_a = ws['A22']
        cell_22_a.value = '내부\n육안검사'
        cell_22_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E22:E25')
        cell_22_e = ws['E22']
        cell_22_e.value = '두께측정'
        cell_22_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A33:A38')
        cell_33_a = ws['A33']
        cell_33_a.value = 'Bundle\n육안검사'
        cell_33_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E26:E34')
        cell_26_e = ws['E26']
        cell_26_e.value = '내부\nNDT'
        cell_26_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A39:A45')
        cell_39_a = ws['A39']
        cell_39_a.value = '외부 육안\n검사'
        cell_39_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E35:E42')
        cell_35_e = ws['E35']
        cell_35_e.value = '추가\n보수작업'
        cell_35_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E43:E45')
        cell_43_e = ws['E43']
        cell_43_e.value = '수압시험'
        cell_43_e.alignment = Alignment(horizontal='center', vertical='center')

        # 테두리 적용
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))

        # Apply the border to all cells in the range
        for row in ws.iter_rows(min_row=1, max_row=num_rows, min_col=1, max_col=num_columns):
            for cell in row:
                cell.border = thin_border
      ########################################################################### 여기부터 수정 필요           
    # =============================================================================================================

        # 장치 번호 입력
        device_number = sheet_data.iloc[0, 0]  # 장치 번호 추정 위치
        ws['B3'].value = device_number  # 장치 번호 셀 할당
        ws['B3'].alignment = Alignment(horizontal='center', vertical='center')

    # =============================================================================================================    

        # 작업 사항 및 검사 사항 사이의 데이터를 처리하기 위한 로직
        worklist_start = None
        content_end = None  # 작업 사항과 검사 사항 사이의 내용을 저장할 위치
        separate_row_found = False  # '띄어쓰기' 행이 발견되었는지 확인

        for i, row in sheet_data.iterrows():
            # '[작업사항]'이 발견되면 시작 위치를 설정합니다.
            if '[작업사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                worklist_start = i

            # '[검사사항]'이 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '[검사사항]' in str(row[0]) and worklist_start is not None:
                content_end = i
                break  # '[검사사항]'이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

            # 띄어쓰기가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif row[0] == '띄어쓰기' and worklist_start is not None:  # row[0]이 빈 문자열인 경우를 확인합니다.
                content_end = i
                empty_row_found = True  # 빈 행이 발견되었음을 표시합니다.
                break  # 빈 행이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 작업 사항과 검사 사항 사이에 데이터가 있고, '[검사사항]' 또는 빈 행이 발견된 경우 데이터 삽입을 시작합니다.
        if worklist_start is not None and content_end is not None:
            worklist_data = sheet_data.iloc[worklist_start + 1:content_end]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(worklist_data[first_column_name], start=9):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

    # =============================================================================================================  
        # 검사 사항과 '띄어쓰기' 사이의 데이터를 처리하기 위한 로직
        inpection_start = None
        content_end1 = None  # 검사 사항과 띄어쓰기 사이의 내용을 저장할 위치

        for i, row in sheet_data.iterrows():
            # '[검사사항]'이 발견되면 시작 위치를 설정합니다.
            if '[검사사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                inpection_start = i

            # '띄어쓰기'가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '띄어쓰기' in str(row[0]) and inpection_start is not None:
                content_end1 = i
                break  # '띄어쓰기'가 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 검사 사항과 '띄어쓰기' 사이에 데이터 삽입을 시작합니다.
        if inpection_start is not None and content_end1 is not None:
            inpection_data = sheet_data.iloc[inpection_start + 1:content_end1]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name1 = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # F9셀부터 데이터 삽입
            for i, item in enumerate(inpection_data[first_column_name1], start=9):
                cell = ws.cell(row=i, column=6)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

        # 워크북을 저장합니다. 파일 이름은 원하는 대로 설정할 수 있습니다.
        output_excel_file = "HEX_Checklist 통합.xlsx"
        saved_file_path = os.path.join(output_folder, output_excel_file)
        wb.save(saved_file_path)
    
    return saved_file_path

# afc checklist 통합본 생성 함수
def afc_checklist_merge(uploaded_file):
    # 파일 로드
    grouped_data = pd.read_excel(uploaded_file, sheet_name=None)

    output_folder = "Checklist_merged_Temp"
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)

    # 새로운 Excel 워크북을 생성합니다.
    wb = openpyxl.Workbook()

    # 기본 제공되는 빈 워크시트 제거
    wb.remove(wb.active)
    
    for sheet_name, sheet_data in grouped_data.items():
        # 새 워크시트를 생성하고 이름을 설정합니다.
        ws = wb.create_sheet(title=sheet_name)

        # Define the number of rows and columns
        num_rows = 45
        num_columns = 8  # Columns A to H

        # Populate the table with placeholder data
        for row in range(1, num_rows + 1):
            for col in range(1, num_columns + 1):
                cell = ws[get_column_letter(col) + str(row)]
                cell.value = None

        # 열 너비 설정
        column_widths = {'A': 14,'B': 35.5,'C': 7.9,'D': 29,'E': 12,'F': 28,'G': 21,'H': 29}

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # 행 높이 설정
        row_heights = {1: 15, 6: 15, 2: 37.8}

        # Apply the specified heights
        for row_num, height in row_heights.items():
            ws.row_dimensions[row_num].height = height

        # Set the height for all other rows
        for row_num in range(1, num_rows + 1):
            if row_num not in row_heights:
                ws.row_dimensions[row_num].height = 30

        # 2행 병합
        ws.merge_cells('A2:H2')

        # 2행 문자 삽입
        cell_2 = ws['A2']
        cell_2.value = "AFC 개방검사 검사표"
        cell_2.font = Font(bold=True, underline='single', size=25)
        cell_2.alignment = Alignment(horizontal='center', vertical='center')

        # 3~5행 병합
        ws.merge_cells('C3:E3')
        ws.merge_cells('C4:E4')
        ws.merge_cells('C5:E5')

        # 3행 문자 삽입
        cell_3_a = ws['A3']
        cell_3_a.value = '장치번호'
        cell_3_a.font = Font(bold=True, size=14)
        cell_3_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_3_f = ws['F3']
        cell_3_f.value = '장치등급'
        cell_3_f.font = Font(bold=True, size=14)
        cell_3_f.alignment = Alignment(horizontal='left', vertical='center')

        # 4행 문자 삽입
        cell_4_a = ws['A4']
        cell_4_a.value = '검사일 :'
        cell_4_a.font = Font(bold=True, size=14)
        cell_4_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_4_f = ws['F4']
        cell_4_f.value = '검사구분'
        cell_4_f.font = Font(bold=True, size=14)
        cell_4_f.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 문자 삽입
        cell_5_a = ws['A5']
        cell_5_a.value = '검사원 :'
        cell_5_a.font = Font(bold=True, size=14)
        cell_5_a.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_c = ws['C5']
        cell_5_c.value = '(서명)'
        cell_5_c.font = Font(bold=True, size=11)
        cell_5_c.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_f = ws['F5']
        cell_5_f.value = '엔지니어 :'
        cell_5_f.font = Font(bold=True, size=14)
        cell_5_f.alignment = Alignment(horizontal='left', vertical='center')

        cell_5_h = ws['H5']
        cell_5_h.value = '(서명)'
        cell_5_h.font = Font(bold=True, size=11)
        cell_5_h.alignment = Alignment(horizontal='left', vertical='center')

        # 5행 아래 굵은 테두리 표시
        bold_bottom_border = Border(bottom=Side(style='thick'))

        for col in range(1, 9):
            ws.cell(row=5, column=col).border = bold_bottom_border

        # 7행 병합 및 문자 삽입
        ws.merge_cells('A7:H7')
        cell_7_a = ws['A7']
        cell_7_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_7_a.value = '1. TA Worklist'
        cell_7_a.font = Font(bold=True, size=12)
        cell_7_a.alignment = Alignment(horizontal='left', vertical='center')

        # 8행 병합 및 문자 삽입
        ws.merge_cells('B8:E8')
        cell_8_a = ws['A8']
        cell_8_a.value = '번호'
        cell_8_a.font = Font(bold=True, size=12)
        cell_8_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('F8:H8')
        cell_8_f = ws['F8']
        cell_8_f.value = '주요 검사 사항'
        cell_8_f.font = Font(bold=True, size=12)
        cell_8_f.alignment = Alignment(horizontal='center', vertical='center')

        cell_8_b = ws['B8']
        cell_8_b.value = '작업 내용'
        cell_8_b.font = Font(bold=True, size=12)
        cell_8_b.alignment = Alignment(horizontal='center', vertical='center')

        # 9~18행 병합 및 문자 삽입
        for i in range(9, 19):
            ws.merge_cells(f'B{i}:E{i}')
            ws.merge_cells(f'F{i}:H{i}')
            cell = ws[f'A{i}']
            cell.value = str(i - 8)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')      

        # 19행 병합 및 문자 삽입
        ws.merge_cells('A19:H19')
        cell_19_a = ws['A19']
        cell_19_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_19_a.value = '2. Damage Mechanism :'
        cell_19_a.font = Font(bold=True, size=12)
        cell_19_a.alignment = Alignment(horizontal='left', vertical='center')

        # 20행 병합 및 문자 삽입
        ws.merge_cells('A20:H20')
        cell_20_a = ws['A20']
        cell_20_a.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        cell_20_a.value = '3. 장치 Type별 일반 검사 항목 (양호, 불량, N/A)'
        cell_20_a.font = Font(bold=True, size=12)
        cell_20_a.alignment = Alignment(horizontal='left', vertical='center')

        # 21행 병합 및 문자 삽입
        cell_21_a = ws['A21']
        cell_21_a.value = '구분'
        cell_21_a.font = Font(bold=True, size=12)
        cell_21_a.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_b = ws['B21']
        cell_21_b.value = '검사항목'
        cell_21_b.font = Font(bold=True, size=12)
        cell_21_b.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('C21:D21')
        cell_21_c = ws['C21']
        cell_21_c.value = '검사 결과'
        cell_21_c.font = Font(bold=True, size=12)
        cell_21_c.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_e = ws['E21']
        cell_21_e.value = '구분'
        cell_21_e.font = Font(bold=True, size=12)
        cell_21_e.alignment = Alignment(horizontal='center', vertical='center')

        cell_21_f = ws['F21']
        cell_21_f.value = '검사항목'
        cell_21_f.font = Font(bold=True, size=12)
        cell_21_f.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('G21:H21')
        cell_21_g = ws['G21']
        cell_21_g.value = '검사 결과'
        cell_21_g.font = Font(bold=True, size=12)
        cell_21_g.alignment = Alignment(horizontal='center', vertical='center')

        # 22행 문자 삽입
        cell_22_b = ws['B22']
        cell_22_b.value = '청소 전 초기상태'
        cell_22_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_22_f= ws['F22']
        cell_22_f.value = 'Shell & Head(Header Box)'
        cell_22_f.alignment = Alignment(horizontal='center', vertical='center')

        # 23행 병합 및 문자 삽입
        cell_23_b = ws['B23']
        cell_23_b.value = 'Sludge, Scale 부착정도'
        cell_23_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_23_f= ws['F23']
        cell_23_f.value = 'Channel & Cover'
        cell_23_f.alignment = Alignment(horizontal='center', vertical='center')

        # 24행 병합 및 문자 삽입
        cell_24_b = ws['B24']
        cell_24_b.value = 'Shell & Head'
        cell_24_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_24_f= ws['F24']
        cell_24_f.value = 'Nozzle & Flange'
        cell_24_f.alignment = Alignment(horizontal='center', vertical='center')

        # 25행 문자 삽입
        cell_25_b = ws['B25']
        cell_25_b.value = 'Channel & Cover'
        cell_25_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_25_f= ws['F25']
        cell_25_f.value = 'IRIS/ECT (본)'
        cell_25_f.alignment = Alignment(horizontal='center', vertical='center')

        # 26행 문자 삽입
        cell_26_b = ws['B26']
        cell_26_b.value = 'Nozzle & Flange'
        cell_26_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_26_f= ws['F26']
        cell_26_f.value = 'Shell/Channel Main 용접심'
        cell_26_f.alignment = Alignment(horizontal='center', vertical='center')

        # 27행 문자 삽입
        cell_27_b = ws['B27']
        cell_27_b.value = 'Floating Head'
        cell_27_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_27_f = ws['F27']
        cell_27_f.value = 'Shell/Channel 부착물 용접심'
        cell_27_f.alignment = Alignment(horizontal='center', vertical='center')

        # 28행 문자 삽입
        cell_28_b = ws['B28']
        cell_28_b.value = 'Header Box'
        cell_28_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_28_f = ws['F28']
        cell_28_f.value = 'Shell/Channel to Nozzle 용접심'
        cell_28_f.alignment = Alignment(horizontal='center', vertical='center')

        # 29행 문자 삽입
        cell_29_b = ws['B29']
        cell_29_b.value = 'Header Box 나사산'
        cell_29_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_29_f= ws['F29']
        cell_29_f.value = 'Nozzle C & B Joint'
        cell_29_f.alignment = Alignment(horizontal='center', vertical='center')

        # 30행 문자 삽입
        cell_30_b = ws['B30']
        cell_30_b.value = 'Plug Bolt 나사산'
        cell_30_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_30_f= ws['F30']
        cell_30_f.value = 'U-Bend RT'
        cell_30_f.alignment = Alignment(horizontal='center', vertical='center')

        # 31행 문자 삽입
        cell_31_b = ws['B31']
        cell_31_b.value = 'Plug Gasket Face'
        cell_31_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_31_f= ws['F31']
        cell_31_f.value = '상부 1열 RT(Fixed Type)'
        cell_31_f.alignment = Alignment(horizontal='center', vertical='center')

        # 32행 문자 삽입
        cell_32_b = ws['B32']
        cell_32_b.value = 'Coating/Lining 상태'
        cell_32_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_32_f= ws['F32']
        cell_32_f.value = 'Strength Welding부 PT'
        cell_32_f.alignment = Alignment(horizontal='center', vertical='center')

        # 33행 문자 삽입
        cell_33_b = ws['B33']
        cell_33_b.value = 'Tube 외부'
        cell_33_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_33_f= ws['F33']
        cell_33_f.value = 'Header Box 용접심'
        cell_33_f.alignment = Alignment(horizontal='center', vertical='center')

        # 34행 문자 삽입
        cell_34_b = ws['B34']
        cell_34_b.value = 'Tube 내부'
        cell_34_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_34_f= ws['F34']
        cell_34_f.value = 'Partition/Stiffener Plate 용접부'
        cell_34_f.alignment = Alignment(horizontal='center', vertical='center')

        # 35행 문자 삽입
        cell_35_b = ws['B35']
        cell_35_b.value = 'Tube Sheet'
        cell_35_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_35_f= ws['F35']
        cell_35_f.value = '용접'
        cell_35_f.alignment = Alignment(horizontal='center', vertical='center')

        # 36행 문자 삽입
        cell_36_b = ws['B36']
        cell_36_b.value = 'Baffle, Tie-Rods'
        cell_36_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_36_f= ws['F36']
        cell_36_f.value = '열처리'
        cell_36_f.alignment = Alignment(horizontal='center', vertical='center')

        # 37행 문자 삽입
        cell_37_b = ws['B37']
        cell_37_b.value = 'Sealing Strip'
        cell_37_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_37_f = ws['F37']
        cell_37_f.value = 'Grinding'
        cell_37_f.alignment = Alignment(horizontal='center', vertical='center')

        # 38행 문자 삽입
        cell_38_b = ws['B38']
        cell_38_b.value = 'Impingement Plate'
        cell_38_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_38_f = ws['F38']
        cell_38_f.value = '보수부위 NDT'
        cell_38_f.alignment = Alignment(horizontal='center', vertical='center')

        # 39행 문자 삽입
        cell_39_b = ws['B39']
        cell_39_b.value = '보온재'
        cell_39_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_39_f = ws['F39']
        cell_39_f.value = '압력시험'
        cell_39_f.alignment = Alignment(horizontal='center', vertical='center')

        # 40행 문자 삽입
        cell_40_b = ws['B40']
        cell_40_b.value = 'Nozzle & Flange'
        cell_40_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_40_f = ws['F40']
        cell_40_f.value = 'Painting'
        cell_40_f.alignment = Alignment(horizontal='center', vertical='center')

        # 41행 문자 삽입
        cell_41_b = ws['B41']
        cell_41_b.value = 'Fittings'
        cell_41_b.alignment = Alignment(horizontal='center', vertical='center')

        # 42행 문자 삽입
        cell_42_b = ws['B42']
        cell_42_b.value = 'Supports'
        cell_42_b.alignment = Alignment(horizontal='center', vertical='center')

        # 43행 문자 삽입
        cell_43_b = ws['B43']
        cell_43_b.value = 'Fire Proofing'
        cell_43_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_43_f = ws['F43']
        cell_43_f.value = '1차'
        cell_43_f.alignment = Alignment(horizontal='center', vertical='center')

        # 44행 문자 삽입
        cell_44_b = ws['B44']
        cell_44_b.value = 'Paint 상태'
        cell_44_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_44_f = ws['F44']
        cell_44_f.value = '2차'
        cell_44_f.alignment = Alignment(horizontal='center', vertical='center')

        # 45행 문자 삽입
        cell_45_b = ws['B45']
        cell_45_b.value = 'Floating Head Anchor bolt\nDe-bolting 유무 점검'
        cell_45_b.alignment = Alignment(horizontal='center', vertical='center')

        cell_45_f = ws['F45']
        cell_45_f.value = '3차'
        cell_45_f.alignment = Alignment(horizontal='center', vertical='center')

        # 세로 병합
        ws.merge_cells('A22:A32')
        cell_22_a = ws['A22']
        cell_22_a.value = '내부\n육안검사'
        cell_22_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E22:E25')
        cell_22_e = ws['E22']
        cell_22_e.value = '두께측정'
        cell_22_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A33:A38')
        cell_33_a = ws['A33']
        cell_33_a.value = 'Bundle\n육안검사'
        cell_33_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E26:E34')
        cell_26_e = ws['E26']
        cell_26_e.value = '내부\nNDT'
        cell_26_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('A39:A45')
        cell_39_a = ws['A39']
        cell_39_a.value = '외부 육안\n검사'
        cell_39_a.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E35:E42')
        cell_35_e = ws['E35']
        cell_35_e.value = '추가\n보수작업'
        cell_35_e.alignment = Alignment(horizontal='center', vertical='center')

        ws.merge_cells('E43:E45')
        cell_43_e = ws['E43']
        cell_43_e.value = '수압시험'
        cell_43_e.alignment = Alignment(horizontal='center', vertical='center')

        # 테두리 적용
        thin_border = Border(left=Side(style='thin'),
                             right=Side(style='thin'),
                             top=Side(style='thin'),
                             bottom=Side(style='thin'))

        # Apply the border to all cells in the range
        for row in ws.iter_rows(min_row=1, max_row=num_rows, min_col=1, max_col=num_columns):
            for cell in row:
                cell.border = thin_border
      ########################################################################### 여기부터 수정 필요           
    # =============================================================================================================

        # 장치 번호 입력
        device_number = sheet_data.iloc[0, 0]  # 장치 번호 추정 위치
        ws['B3'].value = device_number  # 장치 번호 셀 할당
        ws['B3'].alignment = Alignment(horizontal='center', vertical='center')

    # =============================================================================================================    

        # 작업 사항 및 검사 사항 사이의 데이터를 처리하기 위한 로직
        worklist_start = None
        content_end = None  # 작업 사항과 검사 사항 사이의 내용을 저장할 위치
        separate_row_found = False  # '띄어쓰기' 행이 발견되었는지 확인

        for i, row in sheet_data.iterrows():
            # '[작업사항]'이 발견되면 시작 위치를 설정합니다.
            if '[작업사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                worklist_start = i

            # '[검사사항]'이 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '[검사사항]' in str(row[0]) and worklist_start is not None:
                content_end = i
                break  # '[검사사항]'이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

            # 띄어쓰기가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif row[0] == '띄어쓰기' and worklist_start is not None:  # row[0]이 빈 문자열인 경우를 확인합니다.
                content_end = i
                empty_row_found = True  # 빈 행이 발견되었음을 표시합니다.
                break  # 빈 행이 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 작업 사항과 검사 사항 사이에 데이터가 있고, '[검사사항]' 또는 빈 행이 발견된 경우 데이터 삽입을 시작합니다.
        if worklist_start is not None and content_end is not None:
            worklist_data = sheet_data.iloc[worklist_start + 1:content_end]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # B9셀부터 데이터 삽입
            for i, item in enumerate(worklist_data[first_column_name], start=9):
                cell = ws.cell(row=i, column=2)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

    # =============================================================================================================  
        # 검사 사항과 '띄어쓰기' 사이의 데이터를 처리하기 위한 로직
        inpection_start = None
        content_end1 = None  # 검사 사항과 띄어쓰기 사이의 내용을 저장할 위치

        for i, row in sheet_data.iterrows():
            # '[검사사항]'이 발견되면 시작 위치를 설정합니다.
            if '[검사사항]' in str(row[0]):  # row[0]은 첫 번째 열입니다.
                inpection_start = i

            # '띄어쓰기'가 발견되면 해당 위치를 기록하고 반복을 종료합니다.
            elif '띄어쓰기' in str(row[0]) and inpection_start is not None:
                content_end1 = i
                break  # '띄어쓰기'가 발견되면 더 이상 순회할 필요가 없으므로 반복을 종료합니다.

        # 검사 사항과 '띄어쓰기' 사이에 데이터 삽입을 시작합니다.
        if inpection_start is not None and content_end1 is not None:
            inpection_data = sheet_data.iloc[inpection_start + 1:content_end1]  # 시작과 끝 위치 사이의 데이터를 가져옵니다.

            first_column_name1 = sheet_data.columns[0]  # 첫 번째 열의 이름을 가져옵니다.

            # F9셀부터 데이터 삽입
            for i, item in enumerate(inpection_data[first_column_name1], start=9):
                cell = ws.cell(row=i, column=6)

                if isinstance(cell, MergedCell):  # 셀이 병합된 경우
                    # 병합된 셀의 범위를 찾아 병합 해제
                    for range_ in ws.merged_cells.ranges:
                        if cell.coordinate in range_:
                            ws.unmerge_cells(str(range_))
                            break

                    cell = ws.cell(row=i, column=2)  # 병합 해제 후 실제 셀을 다시 가져옴

                cell.value = item  # 실제 셀에 데이터 쓰기

        # 워크북을 저장합니다. 파일 이름은 원하는 대로 설정할 수 있습니다.
        output_excel_file = "AFC_Checklist 통합.xlsx"
        saved_file_path = os.path.join(output_folder, output_excel_file)
        wb.save(saved_file_path)
    
    return saved_file_path

# cdv ta handbook 생성 함수
def create_handbook_cdv(excel_file, process_checklist_title):
    # 엑셀 파일 로드
    workbook = load_workbook(excel_file)
    sheet_names = workbook.sheetnames
    
    # 'Handbook' 폴더 생성 (존재하지 않는 경우)
    output_folder = 'Handbook_Temp'
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)

    # 새 워드 문서 생성
    doc = Document()
    
    # 목차 설정
    doc.add_heading(process_checklist_title, level=1)
    doc.add_heading('목차', level=2)
    doc.add_paragraph('1. CDV, HTR')
    doc.add_paragraph('2. HEX')
    doc.add_paragraph('3. AFC')

    doc.add_page_break()

    # 각 워크시트를 개별 페이지로 변환
    for sheet_name in sheet_names:
        sheet = workbook[sheet_name]

        # 페이지 제목 (장치번호)
        device_number = sheet['B3'].value
        heading = doc.add_heading(f'장치번호 : \n{device_number}', 0)
        run = heading.runs[0]  # Assuming there's only one run in the heading
        run.font.size = Pt(20)  # Set the font size to 20

        # 첫 번째 표 (TA Worklist)
        doc.add_heading('1. TA Worklist', level=1)
        table1 = doc.add_table(rows=1, cols=2)
        table1.style = 'Table Grid'
        hdr_cells1 = table1.rows[0].cells
        hdr_cells1[0].text = '번호'
        hdr_cells1[1].text = '작업 사항'

        # B9에서 B18까지의 셀에서 데이터 가져오기
        items1 = []
        for i in range(9, 19):  # B9에서 B18까지
            value = sheet[f'B{i}'].value
            if value:  # 값이 있는 경우에만 추가
                items1.append(value)

        # '번호' 열에 순서 번호 부여
        for idx, item in enumerate(items1, start=1):
            row = table1.add_row().cells
            row[0].text = str(idx)  # 번호
            row[1].text = item  # 작업 사항

        # 두 번째 표 (Damage Mechanism)
        dm = sheet['A19'].value
        doc.add_heading(f'{dm}', level=1)
        
        # 세 번째 표 (중요 검사 항목)
        doc.add_heading('3. 중요 검사 항목', level=1)
        table3 = doc.add_table(rows=1, cols=2)
        table3.style = 'Table Grid'
        hdr_cells3 = table3.rows[0].cells
        hdr_cells3[0].text = '번호'
        hdr_cells3[1].text = '검사사항'

        # B22에서 B31까지의 셀에서 데이터 가져오기
        items3 = []
        for i in range(22, 32): 
            value = sheet[f'B{i}'].value
            if value:  # 값이 있는 경우에만 추가
                items3.append(value)

        # '번호' 열에 순서 번호 부여
        for idx, item in enumerate(items3, start=1):
            row = table3.add_row().cells
            row[0].text = str(idx)  # 번호
            row[1].text = item  # 검사사항
        
        # 테이블 크기 조정, 가운데 맞춤
        for table in [table1, table3]:  # Assuming table1 and table3 are your table variables
            # Set column widths by adjusting all cells in the column
            column_widths = [Cm(1.19), Cm(14.46)]  # The widths for each column
            for row in table.rows:
                for idx, width in enumerate(column_widths):
                    row.cells[idx].width = width

            # Center-align the text in the first row (header row)
            for cell in table.rows[0].cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 페이지 끝에 페이지 나누기 추가
        doc.add_page_break()
        
        # 문서 저장
        saved_file_path = os.path.join(output_folder, 'handbook.docx')
        doc.save(saved_file_path)
        
    return saved_file_path

# hex, afc ta handbook 생성 함수
def create_handbook_hex_afc(excel_file):
    # 엑셀 파일 로드
    workbook = load_workbook(excel_file)
    sheet_names = workbook.sheetnames
    
    # 'Handbook' 폴더 생성 (존재하지 않는 경우)
    output_folder = 'Handbook_Temp'
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)

    # 새 워드 문서 생성
    doc = Document()
    
    # # 목차 설정
    # doc.add_heading(process_checklist_title, level=1)
    # doc.add_heading('목차', level=2)
    # doc.add_paragraph('1. CDV, HTR')
    # doc.add_paragraph('2. HEX')
    # doc.add_paragraph('3. AFC')
    # doc.add_page_break()

   # 각 워크시트를 개별 페이지로 변환
    for sheet_name in sheet_names:
        sheet = workbook[sheet_name]

        # 페이지 제목 (장치번호)
        device_number = sheet['B3'].value
        heading = doc.add_heading(f'장치번호 : \n{device_number}', 0)
        run = heading.runs[0]  # Assuming there's only one run in the heading
        run.font.size = Pt(20)  # Set the font size to 20

        # 첫 번째 표 (TA Worklist)
        doc.add_heading('1. TA Worklist', level=1)
        table1 = doc.add_table(rows=1, cols=2)
        table1.style = 'Table Grid'
        hdr_cells1 = table1.rows[0].cells
        hdr_cells1[0].text = '번호'
        hdr_cells1[1].text = '작업 사항'

        # B9에서 B18까지의 셀에서 데이터 가져오기
        items1 = []
        for i in range(9, 19):  # B9에서 B18까지
            value = sheet[f'B{i}'].value
            if value:  # 값이 있는 경우에만 추가
                items1.append(value)

        # '번호' 열에 순서 번호 부여
        for idx, item in enumerate(items1, start=1):
            row = table1.add_row().cells
            row[0].text = str(idx)  # 번호
            row[1].text = item  # 작업 사항

        # 두 번째 표 (Damage Mechanism) - 내용이 비어 있으므로 예제에서는 구체적인 내용을 추가하지 않았습니다.
        dm = sheet['A19'].value
        doc.add_heading(f'{dm}', level=1)
        
        # 세 번째 표 (고장 원인 항목)
        doc.add_heading('3. 중요 검사 항목', level=1)
        table3 = doc.add_table(rows=1, cols=2)
        table3.style = 'Table Grid'
        hdr_cells3 = table3.rows[0].cells
        hdr_cells3[0].text = '번호'
        hdr_cells3[1].text = '검사사항'

        # F9에서 F18까지의 셀에서 데이터 가져오기
        items3 = []
        for i in range(9, 19):  # F9에서 F18까지
            value = sheet[f'F{i}'].value
            if value:  # 값이 있는 경우에만 추가
                items3.append(value)

        # '번호' 열에 순서 번호 부여
        for idx, item in enumerate(items3, start=1):
            row = table3.add_row().cells
            row[0].text = str(idx)  # 번호
            row[1].text = item  # 검사사항
        
        # 테이블 크기 조정, 가운데 맞춤
        for table in [table1, table3]:  # Assuming table1 and table3 are your table variables
            # Set column widths by adjusting all cells in the column
            column_widths = [Cm(1.19), Cm(14.46)]  # The widths for each column
            for row in table.rows:
                for idx, width in enumerate(column_widths):
                    row.cells[idx].width = width

            # Center-align the text in the first row (header row)
            for cell in table.rows[0].cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 페이지 끝에 페이지 나누기 추가
        # doc.add_page_break()
        
        # 문서 저장
        saved_file_path = os.path.join(output_folder, 'handbook.docx')
        doc.save(saved_file_path)
        
    return saved_file_path

# handbook 통합 함수
def integrate_docx_files(docx_files):
    
    # 'Handbook' 폴더 생성 (존재하지 않는 경우)
    output_folder = 'Handbook_Temp'
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    else:
        # Clear existing files in the output_folder
        clear_directory(output_folder)
        
    master_doc = Document()

    for i, docx_file in enumerate(docx_files):
        if i != 0:
            master_doc.add_page_break()
        # Load the document from the uploaded file
        sub_doc = Document(io.BytesIO(docx_file.getvalue()))

        # Copy each paragraph and table
        for element in sub_doc.element.body:
            master_doc.element.body.append(element)
            
            
    # 문서 저장
    saved_file_path = os.path.join(output_folder, 'handbook_integrated.docx')
    master_doc.save(saved_file_path)
        
    return saved_file_path
